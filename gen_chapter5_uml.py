# -*- coding: utf-8 -*-
"""
为毕业论文第5章各小节生成UML类图 PNG，并插入现有 thesis docx。
同时生成图源码附件 docx。
"""

import os, zlib, copy, urllib.request, urllib.error, sys, time
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────
# 路径配置
# ──────────────────────────────────────────
DESKTOP     = Path(r"C:\Users\25768\Desktop\专设毕设文件")
THESIS_PATH = DESKTOP / "基于虚幻引擎的动作角色扮演游戏设计与实现.docx"
APPENDIX_PATH = DESKTOP / "基于虚幻引擎的动作角色扮演游戏设计与实现-附件.docx"
UML_DIR     = DESKTOP / "uml"
UML_DIR.mkdir(parents=True, exist_ok=True)

SZ_X4 = Pt(12)
SZ_5  = Pt(10.5)

# ──────────────────────────────────────────
# PlantUML 编码 & 下载
# ──────────────────────────────────────────

def _encode_6bit(b: int) -> str:
    if b < 10:  return chr(ord('0') + b)
    b -= 10
    if b < 26:  return chr(ord('A') + b)
    b -= 26
    if b < 26:  return chr(ord('a') + b)
    b -= 26
    return '-' if b == 0 else '_'

def _append3(b1, b2, b3) -> str:
    return (
        _encode_6bit(b1 >> 2) +
        _encode_6bit(((b1 & 3) << 4) | (b2 >> 4)) +
        _encode_6bit(((b2 & 0xF) << 2) | (b3 >> 6)) +
        _encode_6bit(b3 & 0x3F)
    )

def encode_plantuml(text: str) -> str:
    data = zlib.compress(text.encode('utf-8'), 9)[2:-4]   # raw deflate
    out, i = '', 0
    while i + 2 < len(data):
        out += _append3(data[i], data[i+1], data[i+2]); i += 3
    rem = len(data) - i
    if rem == 1:   out += _append3(data[i], 0, 0)[:2]
    elif rem == 2: out += _append3(data[i], data[i+1], 0)[:3]
    return out

def download_plantuml(puml: str, out_path: Path, retries=3) -> bool:
    encoded = encode_plantuml(puml)
    url = f"https://www.plantuml.com/plantuml/png/{encoded}"
    for attempt in range(retries):
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=30) as resp:
                out_path.write_bytes(resp.read())
            print(f"  ✓ 下载成功: {out_path.name}")
            return True
        except Exception as e:
            print(f"  ✗ 第{attempt+1}次失败: {e}")
            time.sleep(2)
    return False

# ──────────────────────────────────────────
# PlantUML 源码定义（8 张图）
# ──────────────────────────────────────────

SKINPARAM = """\
skinparam classAttributeIconSize 0
skinparam backgroundColor #FFFFFF
skinparam class {
    BackgroundColor #F8F9FA
    BorderColor #495057
    ArrowColor #495057
    FontSize 11
}
skinparam stereotype {
    CBackgroundColor #D0E8FF
    CBorderColor #3399CC
}
skinparam note {
    BackgroundColor #FFFDE7
    BorderColor #F9A825
}
"""

DIAGRAMS = [
    # (文件名前缀, 图题, 搜索标题关键词, PlantUML源码)

    ("5_1_architecture", "图5-1\u2002系统分层架构类图", "整体架构与分层",
     f"""@startuml
{SKINPARAM}
title 第5章 5.1 系统分层架构

package "UGameInstance 层（跨关卡持久）" #E8F4FD {{
  class ARPGEventManager <<GameInstanceSubsystem>> {{
    - listeners : TMap<FName, TArray<FARPGObjectListener>>
    + AddListener(EventName, Object, Callback)
    + RemoveListener(EventName, Object)
    + Broadcast(EventName, Payload)
  }}
  class ARPGSaveGameManager <<GameInstanceSubsystem>> {{
    - currentData : FARPGPlayerData
    + AsyncSave(SlotName)
    + LoadFromSlot(SlotName) : FARPGPlayerData
    + GetPlayerData() : FARPGPlayerData
  }}
  class CharacterManager <<GameInstanceSubsystem>> {{
    - characters : TArray<AActor*>
    + Register(Character)
    + Unregister(Character)
    + GetAllCharacters() : TArray<AActor*>
  }}
  class UWidgetManager <<GameInstanceSubsystem>> {{
    - widgetMap : TMap<FName, UBaseWidget*>
    + ShowWidget(ID)
    + HideWidget(ID)
    + GetWidget(ID) : UBaseWidget*
  }}
}}

package "UWorld 层（关卡内有效）" #E8FDE8 {{
  class PoolSubsystem <<WorldSubsystem>> {{
    - pools : TMap<UClass*, FActorPool>
    + PrewarmPool(Class, Count)
    + RequestActor(Class) : AActor*
    + ReleaseActor(Actor)
  }}
}}

package "业务层（Game Layer）" #FDF3E8 {{
  class InGameMode
  class AInGameCharacter
  class AInGameAICharacter
  class ULobbyWidget
}}

AInGameCharacter    ..> ARPGEventManager  : 广播/订阅
AInGameAICharacter  ..> ARPGEventManager  : 广播/订阅
AInGameCharacter    ..> PoolSubsystem     : 借还 Actor
InGameMode          ..> CharacterManager  : 注册角色
ULobbyWidget        ..> ARPGSaveGameManager : 读写存档
ULobbyWidget        ..> UWidgetManager   : 管理子 Widget
@enduml"""),

    ("5_2_1_character_select", "图5-2\u2002角色选择模块类图", "角色选择实现",
     f"""@startuml
{SKINPARAM}
title 第5章 5.2.1 角色选择模块

class ULobbyWidget {{
  + CharacterSelectPage : UCharacterSelect*
  + MapSelectPage : UMapAndModeSelect*
  + NativeOnInitialized()
}}

class UCharacterSelect {{
  - selectedRowHandle : FDataTableRowHandle
  + RefreshCharacterList()
  + OnCharacterCardClicked(Handle)
  + ShowMoneyDisplay(Amount : int32)
}}

class UMapAndModeSelect {{
  + OnMapSelected(MapID : FName)
}}

class ARPGSaveGameManager <<GameInstanceSubsystem>> {{
  + GetPlayerData() : FARPGPlayerData
  + GetTotalMoney() : int32
}}

struct FARPGPlayerData {{
  + PlayerUserName : FString
  + PlayerTotalMoney : int32
  + HealthMod : float
  + AttackPowerMod : float
  + DefensePowerMod : float
}}

class "CharacterTable\\n[DataTable]" as CT <<DataTable>> {{
  + CharacterID : FName
  + MeshAsset
  + AnimBPClass
  + BaseHealth / BaseAttack / BaseDefense : float
}}

ULobbyWidget *-- UCharacterSelect
ULobbyWidget *-- UMapAndModeSelect
UCharacterSelect --> ARPGSaveGameManager : GetPlayerData()
UCharacterSelect --> CT              : 读取角色行列表
ARPGSaveGameManager ..> FARPGPlayerData : returns
@enduml"""),

    ("5_2_2_shop", "图5-3\u2002商店模块类图", "商店实现",
     f"""@startuml
{SKINPARAM}
title 第5章 5.2.2 商店模块

class UWidgetManager <<GameInstanceSubsystem>> {{
  + InitFromTable(ConfigTable)
  + ShowWidget(ID)
  + HideWidget(ID)
  + GetWidget(ID) : UBaseWidget*
}}

abstract class UBaseWidget {{
  + Show()
  + Hide()
  + OnInit()
}}

class UARPGLeaf_ShopItem {{
  - ItemName   : FText
  - ItemIcon   : UTexture2D*
  - ItemPrice  : int32
  + BindData(Row : FShopItemDataRow)
  + OnBuyButtonClicked()
}}

struct FShopItemDataRow {{
  + ItemID   : FName
  + ItemName : FText
  + Price    : int32
  + EffectGE : TSubclassOf<UGameplayEffect>
}}

class ARPGSaveGameManager <<GameInstanceSubsystem>> {{
  + DeductMoney(Amount) : bool
  + AsyncSave(SlotName)
}}

class ARPGEventManager <<GameInstanceSubsystem>> {{
  + Broadcast(ARPG_MoneyUpdated, Payload)
}}

UBaseWidget <|-- UARPGLeaf_ShopItem
UWidgetManager "1" *-- "*" UBaseWidget
UARPGLeaf_ShopItem --> FShopItemDataRow : binds
UARPGLeaf_ShopItem --> ARPGSaveGameManager : DeductMoney (Server RPC)
UARPGLeaf_ShopItem --> ARPGEventManager   : Broadcast 购买结果
@enduml"""),

    ("5_3_1_player_character", "图5-4\u2002玩家角色类图", "玩家角色实现",
     f"""@startuml
{SKINPARAM}
title 第5章 5.3.1 玩家角色

interface IAbilitySystemInterface {{
  + GetAbilitySystemComponent() : UAbilitySystemComponent*
}}

abstract class ARPGBaseCharacter {{
  # DamageComponent  : UCharacterDamageComponent*
  # AttackComponent  : UCharacterAttackComponent*
  # HealthManager    : UCharacterHealthManager*
  # StaminaManager   : UStaminaManagerComponent*
  # ShieldComponent  : UCharacterShieldComponent*
  # RollComponent    : URollComponent*
  # SprintComponent  : USprintComponent*
  + GetGenericTeamId() : FGenericTeamId
}}

class AInGameCharacter {{
  - ASC             : UAbilitySystemComponent*
  - AttributeSet    : UInGameCharacterAttributeSet*
  - TargetLockCfg   : FTargetLockConfig
  - CurrentTarget   : AActor*
  + GetAbilitySystemComponent()
  + TryLockTarget()
  + ClearLockTarget()
  - SetupEnhancedInput()
}}

struct FTargetLockConfig {{
  + DetectRadius      : float = 1000.0
  + FOVAngle          : float = 60.0
  + CameraFollowSpeed : float = 10.0
}}

class UInGameCharacterAttributeSet <<AttributeSet>> {{
  + Health / MaxHealth       : FGameplayAttributeData
  + AttackPower              : FGameplayAttributeData
  + DefensePower             : FGameplayAttributeData
  + DamageReduction          : FGameplayAttributeData
  + Stamina / MaxStamina     : FGameplayAttributeData
  + AttackStaminaCost        : FGameplayAttributeData
  + DefenseStaminaCost       : FGameplayAttributeData
  + CharacterMoney           : FGameplayAttributeData
}}

ARPGBaseCharacter     <|-- AInGameCharacter
IAbilitySystemInterface <|.. AInGameCharacter
AInGameCharacter *-- FTargetLockConfig
AInGameCharacter *-- UInGameCharacterAttributeSet
@enduml"""),

    ("5_3_2_ai_character", "图5-5\u2002AI角色类图", "AI角色实现",
     f"""@startuml
{SKINPARAM}
title 第5章 5.3.2 AI 角色

abstract class ARPGBaseCharacter

class AInGameAICharacter {{
  - Manifest             : FAIManifest
  - PerceptionComp       : UAIPerceptionComponent*
  + DefaultAttackRange   : float = 200.0
  + AttackCooldown       : float = 1.0
  + InitFromManifest(Manifest)
}}

struct FAIManifest {{
  + BehaviorTreeAsset    : UBehaviorTree*
  + BlackboardAsset      : UBlackboardData*
  + PatrolPoints         : TArray<AActor*>
  + BaseAttributeRow     : FDataTableRowHandle
}}

class InGameAIController <<AIController>> {{
  - CurrentState         : EAIControlState
  - AttackCooldownTimer  : float
  + OnPossess(Pawn)
  + UpdateAIState()
  + SetBlackboardTarget(Target : AActor*)
}}

enum EAIControlState {{
  Patrol
  Attack
  Roll
  Defense
  Dodge
  Heal
}}

class BTService_UpdateAIState <<BTService>> {{
  + StartDodgeStaminaPercent : float = 0.2
  + HealingHealthPercent     : float = 0.3
  + StartDefenseDistance     : float = 10.0
  + TickNode(OwnerComp, DeltaSeconds)
}}

ARPGBaseCharacter    <|-- AInGameAICharacter
AInGameAICharacter   *-- FAIManifest
InGameAIController   --> AInGameAICharacter  : Possess
InGameAIController    o-- EAIControlState    : maintains
BTService_UpdateAIState --> InGameAIController : updates state
@enduml"""),

    ("5_3_3_weapon", "图5-6\u2002武器系统类图", "武器系统实现",
     f"""@startuml
{SKINPARAM}
title 第5章 5.3.3 武器系统

abstract class AARPGBaseWeapon <<AActor>> {{
  - Manifest           : FWeaponManifest
  + Activate(ActivateType : EEquipmentActivateType)
  + Deactivate()
  + GetWeaponType() : EWeaponType
  + ApplyDamageGE(Target : AActor*)
}}

struct FWeaponManifest {{
  + WeaponType         : EWeaponType
  + DamageMultiplier   : float
  + LightAttackMontage : UAnimMontage*
  + HeavyAttackMontage : UAnimMontage*
  + JumpAttackMontage  : UAnimMontage*
  + DefenseMontage     : UAnimMontage*
  + AttackGE           : TSubclassOf<UGameplayEffect>
}}

enum EWeaponType {{
  Unarmeds
  Sword
  Shield
}}

enum EEquipmentActivateType {{
  LightAttack
  HeavyAttack
  JumpAttack
  NormalDefense
  BounceBack
}}

class UEquipmentComponent <<UActorComponent>> {{
  - MainHandWeapon     : AARPGBaseWeapon*
  - OffHandWeapon      : AARPGBaseWeapon*
  + EquipWeapon(Weapon, Slot)
  + UnequipWeapon(Slot)
  + ActivateEquipment(Type : EEquipmentActivateType)
}}

AARPGBaseWeapon *-- FWeaponManifest
AARPGBaseWeapon --> EWeaponType
AARPGBaseWeapon --> EEquipmentActivateType
UEquipmentComponent "1" o-- "0..2" AARPGBaseWeapon
@enduml"""),

    ("5_4_ui_system", "图5-7\u2002UI系统类图", "UI系统实现",
     f"""@startuml
{SKINPARAM}
title 第5章 5.4 UI 系统（组合模式）

abstract class UBaseWidget <<UUserWidget>> {{
  + Show()
  + Hide()
  + OnInit()
}}

abstract class UARPGCompositeBase {{
  - Children : TArray<UBaseWidget*>
  + AddChild(Widget)
  + RemoveChild(Widget)
  + Refresh()
}}

class UARPGLeaf {{
  + BindData(DataObject)
  + Refresh()
}}

class UWidgetManager <<GameInstanceSubsystem>> {{
  - widgetRegistry : TMap<FName, UBaseWidget*>
  + ShowWidget(ID)
  + HideWidget(ID)
  + GetWidget(ID) : UBaseWidget*
  + InitFromTable(ConfigTable)
}}

class ULobbyWidget <<Composite>> {{
  + CharacterSelectPage : UCharacterSelect*
  + MapSelectPage       : UMapAndModeSelect*
}}

class UTargetLockWidget <<Leaf>> {{
  - CurrentTarget : AActor*
  + OnTargetChanged(Target)
  + UpdateLockIndicatorPosition()
}}

class UMiniMapWidget <<Leaf>> {{
  - CaptureComp : USceneCaptureComponent2D*
  + UpdatePlayerIcon(Loc)
  + UpdateEnemyIcons(Locs)
}}

class ULoadingScreenWidget <<Leaf>> {{
  - Progress : float
  + OnLevelStreamingStateChanged(State)
  + UpdateProgressBar(Percent)
}}

UBaseWidget            <|-- UARPGCompositeBase
UBaseWidget            <|-- UARPGLeaf
UARPGCompositeBase     <|-- ULobbyWidget
UBaseWidget            <|-- UTargetLockWidget
UBaseWidget            <|-- UMiniMapWidget
UBaseWidget            <|-- ULoadingScreenWidget
UWidgetManager         "1" *-- "*" UBaseWidget : manages
UARPGCompositeBase     *-- "*" UBaseWidget : children
@enduml"""),

    ("5_5_data_driven", "图5-8\u2002数据驱动DataTable类图", "数据驱动实现",
     f"""@startuml
{SKINPARAM}
title 第5章 5.5 数据驱动（DataTable 结构）

class CharacterTable <<DataTable>> {{
  + CharacterID   : FName
  + MeshAsset     : TSoftObjectPtr<USkeletalMesh>
  + AnimBPClass   : TSubclassOf<UAnimInstance>
  + BaseHealth    : float
  + BaseAttack    : float
  + BaseDefense   : float
  + BaseStamina   : float
}}

class WeaponConfigTable <<DataTable>> {{
  + WeaponID       : FName
  + WeaponType     : EWeaponType
  + MeshAsset      : TSoftObjectPtr<USkeletalMesh>
  + DamageMultiplier : float
}}

class LevelConfigTable <<DataTable>> {{
  + MapID          : FName
  + LevelToOpen    : FName
  + SpawnPoints    : TArray<FSpawnPointData>
  + EscapePoints   : TArray<FEscapePointData>
}}

struct FSpawnPointData {{
  + SelectWeight   : float
  + bIsSelfCamp    : bool
  + SpawnLocation  : FVector
}}

struct FEscapePointData {{
  + bNeedEscapeItem   : bool
  + EscapeWaitingTime : float
  + Location          : FVector
}}

class ShopItemDataTable <<DataTable>> {{
  + ItemID    : FName
  + ItemName  : FText
  + Price     : int32
  + EffectGE  : TSubclassOf<UGameplayEffect>
}}

class WidgetConfigTable <<DataTable>> {{
  + WidgetID    : FName
  + WidgetClass : TSubclassOf<UBaseWidget>
  + ZOrder      : int32
  + bAutoCreate : bool
}}

LevelConfigTable "1" *-- "*" FSpawnPointData
LevelConfigTable "1" *-- "*" FEscapePointData
@enduml"""),
]

# ──────────────────────────────────────────
# 1. 下载所有 PNG
# ──────────────────────────────────────────
print("=" * 55)
print("第一步：下载 PlantUML PNG 图片")
print("=" * 55)

downloaded = {}   # prefix -> Path
failed     = []

for prefix, caption, _, puml in DIAGRAMS:
    out = UML_DIR / f"{prefix}.png"
    print(f"\n正在处理：{caption}")
    ok = download_plantuml(puml, out)
    if ok:
        downloaded[prefix] = out
    else:
        failed.append(prefix)
        print(f"  警告：{prefix} 下载失败，将在 docx 中插入占位文本")

print(f"\n下载完成：{len(downloaded)}/{len(DIAGRAMS)} 成功")

# ──────────────────────────────────────────
# 2. 辅助：在 docx 中查找包含关键词的段落
# ──────────────────────────────────────────

def find_heading_para(doc, keyword):
    """返回第一个包含 keyword 的段落，否则返回 None"""
    for para in doc.paragraphs:
        if keyword in para.text:
            return para
    return None

def set_font_run(run, cn="宋体", en="Times New Roman", size=SZ_5, bold=False):
    run.bold = bold
    run.font.size = size
    run.font.name = en
    # Ensure rPr element exists (python-docx may be lazy)
    rPr = run._element.get_or_add_rPr()
    fonts_el = rPr.find(qn('w:rFonts'))
    if fonts_el is None:
        fonts_el = OxmlElement('w:rFonts')
        rPr.insert(0, fonts_el)
    fonts_el.set(qn('w:eastAsia'), cn)

def build_image_element(doc_ctx, img_path: Path):
    """生成图片段落 XML 元素（居中）"""
    from docx import Document as D2
    tmp = D2()
    p = tmp.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(4.8))
    return copy.deepcopy(p._element)

def build_caption_element(caption_text):
    """生成图题段落 XML 元素（五号宋体，居中）"""
    p_xml = OxmlElement('w:p')
    pPr   = OxmlElement('w:pPr')
    jc    = OxmlElement('w:jc');   jc.set(qn('w:val'), 'center')
    sp_bef = OxmlElement('w:spacing')
    sp_bef.set(qn('w:before'), '0')
    sp_bef.set(qn('w:after'),  '120')
    pPr.append(jc); pPr.append(sp_bef)
    p_xml.append(pPr)

    r_xml = OxmlElement('w:r')
    rPr   = OxmlElement('w:rPr')
    sz    = OxmlElement('w:sz');     sz.set(qn('w:val'), '21')   # 10.5pt
    szCs  = OxmlElement('w:szCs');  szCs.set(qn('w:val'), '21')
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'),    'Times New Roman')
    fonts.set(qn('w:eastAsia'), '宋体')
    rPr.append(fonts); rPr.append(sz); rPr.append(szCs)
    r_xml.append(rPr)
    t = OxmlElement('w:t'); t.text = caption_text
    r_xml.append(t)
    p_xml.append(r_xml)
    return p_xml

def build_placeholder_element(text):
    """若图片下载失败，插入一段灰色占位文字"""
    p_xml = OxmlElement('w:p')
    pPr   = OxmlElement('w:pPr')
    jc    = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
    pPr.append(jc); p_xml.append(pPr)
    r_xml = OxmlElement('w:r')
    rPr   = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), 'FF0000')
    rPr.append(color); r_xml.append(rPr)
    t = OxmlElement('w:t'); t.text = f"[图片未下载：{text}]"
    r_xml.append(t); p_xml.append(r_xml)
    return p_xml

# ──────────────────────────────────────────
# 3. 将图片插入 thesis docx
# ──────────────────────────────────────────
print("\n" + "=" * 55)
print("第二步：将图片插入论文 docx")
print("=" * 55)

if not THESIS_PATH.exists():
    print(f"错误：找不到论文文件 {THESIS_PATH}")
    sys.exit(1)

doc = Document(str(THESIS_PATH))

# 映射: 搜索关键词 -> (图片前缀, 图题文字)
insertions = [
    ("整体架构与分层",   "5_1_architecture",       "图5-1\u2002系统分层架构类图"),
    ("角色选择实现",     "5_2_1_character_select", "图5-2\u2002角色选择模块类图"),
    ("商店实现",         "5_2_2_shop",             "图5-3\u2002商店模块类图"),
    ("玩家角色实现",     "5_3_1_player_character", "图5-4\u2002玩家角色类图"),
    ("AI角色实现",       "5_3_2_ai_character",     "图5-5\u2002AI角色类图"),
    ("武器系统实现",     "5_3_3_weapon",           "图5-6\u2002武器系统类图"),
    ("UI系统实现",       "5_4_ui_system",          "图5-7\u2002UI系统类图"),
    ("数据驱动实现",     "5_5_data_driven",        "图5-8\u2002数据驱动DataTable类图"),
]

inserted_count = 0
for keyword, prefix, caption in insertions:
    heading_para = find_heading_para(doc, keyword)
    if heading_para is None:
        print(f"  [!] 未找到包含 [{keyword}] 的段落，跳过")
        continue

    # 构建图片 & 图题元素
    img_path = UML_DIR / f"{prefix}.png"
    if img_path.exists() and prefix in downloaded:
        img_elem = build_image_element(doc, img_path)
        cap_elem = build_caption_element(caption)
        # 在 heading 之后插入：先插 caption，再插 image（addnext 每次插到紧随其后）
        heading_para._element.addnext(cap_elem)
        heading_para._element.addnext(img_elem)
    else:
        ph_elem  = build_placeholder_element(f"{prefix}.png")
        cap_elem = build_caption_element(caption)
        heading_para._element.addnext(cap_elem)
        heading_para._element.addnext(ph_elem)

    print(f"  + 已在 [{keyword}] 后插入 {caption}")
    inserted_count += 1

THESIS_OUT = THESIS_PATH.with_stem(THESIS_PATH.stem + "_v2")
doc.save(str(THESIS_OUT))
print(f"\n已保存更新后的论文：{THESIS_OUT}")
print(f"（提示：如需覆盖原文件，请先关闭 Word，然后将 _v2 版本重命名）")
print(f"共插入图片段落：{inserted_count} 处")

# ──────────────────────────────────────────
# 4. 生成附件 docx（图源码）
# ──────────────────────────────────────────
print("\n" + "=" * 55)
print("第三步：生成图源码附件 docx")
print("=" * 55)

app = Document()
for section in app.sections:
    section.top_margin    = Mm(30)
    section.bottom_margin = Mm(25)
    section.left_margin   = Mm(30)
    section.right_margin  = Mm(20)

# 封面标题
p = app.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run("基于虚幻引擎的动作角色扮演游戏设计与实现")
set_font_run(r, cn="黑体", size=Pt(16), bold=True)

p2 = app.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(6)
r2 = p2.add_run("附件：第五章UML类图源码")
set_font_run(r2, cn="黑体", size=Pt(14), bold=True)

app.add_page_break()

for prefix, caption, section_title, puml in DIAGRAMS:
    # 小节标题
    ph = app.add_paragraph()
    ph.paragraph_format.space_before = Pt(12)
    ph.paragraph_format.space_after  = Pt(6)
    rh = ph.add_run(f"{caption}（{section_title}）")
    set_font_run(rh, cn="黑体", size=Pt(14), bold=True)

    # 如果图片下载成功，先展示缩略图
    img_path = UML_DIR / f"{prefix}.png"
    if img_path.exists():
        pi = app.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi.add_run().add_picture(str(img_path), width=Inches(5.5))
        pc = app.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rpc = pc.add_run(caption)
        set_font_run(rpc, cn="宋体", size=SZ_5)

    # PlantUML 源码
    src_title = app.add_paragraph()
    rs = src_title.add_run("PlantUML 源码：")
    set_font_run(rs, cn="黑体", size=Pt(11), bold=True)

    for line in puml.strip().split('\n'):
        pc2 = app.add_paragraph()
        pc2.paragraph_format.left_indent  = Pt(24)
        pc2.paragraph_format.space_before = Pt(0)
        pc2.paragraph_format.space_after  = Pt(0)
        pc2.paragraph_format.line_spacing = Pt(16)
        rcode = pc2.add_run(line)
        rcode.font.name = 'Courier New'
        rcode.font.size = Pt(9)
        # 灰色底纹
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pc2._p.pPr.append(shd)

    app.add_paragraph()  # 间距

app.save(str(APPENDIX_PATH))
print(f"已保存附件 docx：{APPENDIX_PATH}")

print("\n" + "=" * 55)
print("全部完成！")
print(f"  论文：{THESIS_PATH}")
print(f"  附件：{APPENDIX_PATH}")
print(f"  PNG 图片目录：{UML_DIR}")
print("=" * 55)
