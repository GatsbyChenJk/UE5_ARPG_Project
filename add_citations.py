# -*- coding: utf-8 -*-
"""
在 _v4_fixed.docx 正文中，根据语义在对应句末插入参考文献上标 [n]。
直接修改原文件，不另存新版本。
"""
import re, pathlib
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

DESKTOP = pathlib.Path.home() / "Desktop" / "\u4e13\u8bbe\u6bd5\u8bbe\u6587\u4ef6"
SRC = DESKTOP / "\u57fa\u4e8e\u865a\u5e7b\u5f15\u64ce\u7684\u52a8\u4f5c\u89d2\u8272\u626e\u6f14\u6e38\u620f\u8bbe\u8ba1\u4e0e\u5b9e\u73b0_v4_fixed.docx"
DST = SRC  # 直接覆盖

def add_superscript(para, label: str):
    """在段落末尾追加上标 [n]，Times New Roman 小五。"""
    run = para.add_run(label)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    # 设置上标
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign)
    # 确保字体
    fonts_el = rPr.find(qn('w:rFonts'))
    if fonts_el is None:
        fonts_el = OxmlElement('w:rFonts')
        rPr.insert(0, fonts_el)
    fonts_el.set(qn('w:ascii'), 'Times New Roman')
    fonts_el.set(qn('w:hAnsi'), 'Times New Roman')

# 引用规则：(关键词列表, 引用标签, 只插入一次)
# 每条规则匹配段落文本，命中则在该段末尾插入上标
# already_inserted 防止同一引用重复插入
CITATION_RULES = [
    # [3] Millington — para 29
    (["Millington", "Funge", "Artificial Intelligence for Games"], "[3]"),
    # [4] 王鑫 UE5渲染 — para 26 Nanite/Lumen
    (["Nanite\u865a\u62df\u5316\u5fae\u591a\u8fb9\u5f62", "Lumen\u5168\u5c40\u5149\u7167"], "[4]"),
    # [2] Game Engine Architecture — para 44
    (["UE5\uff09\u91c7\u7528\u5206\u5c42\u67b6\u6784", "\u5f15\u64ce\u6838\u5fc3\u5c42"], "[2]"),
    # [13] 赵宇飞 UE5入门 — para 46 UObject基类
    (["UObject \u662f\u6240\u6709\u7c7b\u7684\u7ec8\u6781\u57fa\u7c7b", "\u53cd\u5c04\u3001\u5783\u573e\u56de\u6536"], "[13]"),
    # [14] Sherrod — para 48 PrimitiveComponent渲染与物理
    (["UPrimitiveComponent\u63d0\u4f9b\u6e32\u67d3\u4e0e"], "[14]"),
    # [7] Epic GAS文档 — para 55 GAS框架介绍
    (["GAS\uff09\u662fUE5\u63d0\u4f9b\u7684\u4e00\u5957", "AbilitySystemComponent\u3001AttributeSet"], "[7]"),
    # [5] 陈佳明 行为树 — para 60 行为树介绍
    (["Selector/Sequence/Parallel", "\u88c5\u9970\u8282\u70b9\uff08Decorator\uff09"], "[5]"),
    # [11] Abdelkhalek — para 61 AIPerception
    (["AIPerceptionComponent\u4e3aAI Actor\u63d0\u4f9b\u611f\u77e5\u80fd\u529b", "OnTargetPerceptionUpdated"], "[11]"),
    # [8] Epic 网络同步 — para 65 UDP专用服务器
    (["\u4e13\u7528\u670d\u52a1\u5668\uff08Dedicated Server\uff09", "UDP\u534f\u8bae\uff0c\u800c\u539f\u751fUDP"], "[8]"),
    # [10] 吴振宇 多人网络 — para 67 RPC
    (["\u8fdc\u7a0b\u8fc7\u7a0b\u8c03\u7528\uff08RPC\uff09\uff0c\u901a\u8fc7\u53cd\u5c04\u6846\u67b6"], "[10]"),
    # [9] 刘洋 发布订阅 — para 119 事件总线
    (["ARPGEventManager", "\u53d1\u5e03-\u8ba2\u9605"], "[9]"),
    # [6] 黄俊峰 对象池 — para 124 PoolSubsystem
    (["PoolSubsystem\u4f5c\u4e3aWorldSubsystem", "FActorPool"], "[6]"),
    # [12] 程林 数据驱动 — para 74 DataTable
    (["\u6570\u636e\u9a71\u52a8\u8bbe\u8ba1\uff08Data-Driven Design\uff09", "DataTable\u4ee5\u7ed3\u6784\u4f53"], "[12]"),
    # [15] 李明阳 GAS MOBA — para 92 可扩展性需求
    (["\u65b0\u589e\u6280\u80fd\u6216\u72b6\u6001\u6548\u679c\u4ec5\u9700\u6269\u5c55GameplayAbility\u4e0eGameplayEffect"], "[15]"),
    # [1] 游戏编程模式 — para 103 组件化设计
    (["ARPGBaseCharacter\u4f5c\u4e3a\u6240\u6709\u89d2\u8272", "\u7ec4\u4ef6\u5316\u8bbe\u8ba1\u5c06\u6218\u6597\u76f8\u5173\u529f\u80fd\u62c6\u5206"], "[1]"),
]

doc = Document(str(SRC))
inserted = {rule[1]: False for rule in CITATION_RULES}

# 只处理正文（参考文献之前）
in_ref = False
count = 0
for para in doc.paragraphs:
    t = para.text.strip()
    if t in ("\u53c2\u8003\u6587\u732e", "References"):
        in_ref = True
    if in_ref:
        continue
    if not t or not para.runs:
        continue

    for keywords, label in CITATION_RULES:
        if inserted[label]:
            continue
        # 检查所有关键词是否都在段落中（支持正则）
        matched = all(re.search(kw, t) for kw in keywords)
        if matched:
            add_superscript(para, label)
            inserted[label] = True
            count += 1
            print(f"  {label} -> para: {t[:60]}")
            break  # 每段只插一个引用

doc.save(str(DST))
print(f"\nTotal citations inserted: {count}")
print(f"Saved: {DST.name}")
not_inserted = [label for label, done in inserted.items() if not done]
if not_inserted:
    print(f"Not inserted: {not_inserted}")
