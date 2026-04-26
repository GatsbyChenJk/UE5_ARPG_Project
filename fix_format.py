# -*- coding: utf-8 -*-
"""
修复 _v4b.docx（或 _v4.docx）：
1. 所有 run 的英文字体改为 Times New Roman
2. 所有段落行距改为 1.5 倍
3. 参考文献序号后的空格去除（如 "[1] 作者" → "[1]作者"）
"""
import re, pathlib
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DESKTOP = pathlib.Path.home() / "Desktop" / "\u4e13\u8bbe\u6bd5\u8bbe\u6587\u4ef6"

# 优先用 v4b，没有则用 v4
SRC = DESKTOP / "\u57fa\u4e8e\u865a\u5e7b\u5f15\u64ce\u7684\u52a8\u4f5c\u89d2\u8272\u626e\u6f14\u6e38\u620f\u8bbe\u8ba1\u4e0e\u5b9e\u73b0_v4.docx"
DST = SRC.with_stem(SRC.stem.rstrip('b') + "_fixed")

EN_FONT = "Times New Roman"
# 跳过代码段（Courier New / Consolas）
CODE_FONTS = {"Courier New", "Consolas", "Courier", "Lucida Console"}

def fix_run_font(run):
    """将 run 的英文字体设为 Times New Roman（跳过代码字体）"""
    cur = run.font.name
    if cur in CODE_FONTS:
        return
    # 设置 ASCII / hAnsi 字体
    run.font.name = EN_FONT
    # 确保 w:rFonts 存在并设置 w:ascii / w:hAnsi
    rPr = run._element.get_or_add_rPr()
    fonts_el = rPr.find(qn('w:rFonts'))
    if fonts_el is None:
        fonts_el = OxmlElement('w:rFonts')
        rPr.insert(0, fonts_el)
    # 只在未设置或非 Times New Roman 时覆盖
    for attr in (qn('w:ascii'), qn('w:hAnsi')):
        if fonts_el.get(attr) not in (EN_FONT, None) or fonts_el.get(attr) is None:
            fonts_el.set(attr, EN_FONT)

def fix_para_line_spacing(para):
    """将段落行距设为 1.5 倍（跳过已经是 1.5 倍的）"""
    pf = para.paragraph_format
    # 检查是否已经是 1.5 倍
    if (pf.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE):
        return
    # 跳过固定行距（图题、表格内容等可能用固定行距）
    if pf.line_spacing_rule == WD_LINE_SPACING.EXACTLY:
        return
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = None  # 清除具体数值，让规则生效

# 参考文献序号模式：[数字] 后跟空格
REF_PATTERN = re.compile(r'^(\[\d+\])\s+')

def fix_ref_space(para):
    """去除参考文献序号后的空格"""
    if not para.runs:
        return
    first_run = para.runs[0]
    m = REF_PATTERN.match(first_run.text)
    if m:
        first_run.text = first_run.text[len(m.group(0)):].lstrip()
        # 在最前面插入序号（无空格）
        first_run.text = m.group(1) + first_run.text

print(f"Source: {SRC.name}")
doc = Document(str(SRC))

ref_section = False
fixed_font = 0
fixed_spacing = 0
fixed_ref = 0

for para in doc.paragraphs:
    text = para.text.strip()

    # 检测参考文献区域
    if text in ("\u53c2\u8003\u6587\u732e", "References"):
        ref_section = True

    # 1. 修复英文字体
    for run in para.runs:
        old = run.font.name
        fix_run_font(run)
        if run.font.name != old:
            fixed_font += 1

    # 2. 修复行距（跳过表格内段落，表格段落通过 doc.tables 处理）
    fix_para_line_spacing(para)
    fixed_spacing += 1

    # 3. 参考文献序号空格
    if ref_section and REF_PATTERN.match(para.runs[0].text if para.runs else ""):
        fix_ref_space(para)
        fixed_ref += 1

# 表格内段落也修复字体和行距
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    fix_run_font(run)
                # 表格内不强制改行距，保持原样

doc.save(str(DST))
print(f"Done.")
print(f"  Font runs fixed : {fixed_font}")
print(f"  Paragraphs processed for spacing: {fixed_spacing}")
print(f"  Ref entries fixed: {fixed_ref}")
print(f"  Saved: {DST.name}")
