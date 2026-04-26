# -*- coding: utf-8 -*-
import zlib, urllib.request, pathlib, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DESKTOP = pathlib.Path.home() / "Desktop" / "\u4e13\u8bbe\u6bd5\u8bbe\u6587\u4ef6"
UML_DIR = DESKTOP / "uml"
SRC = DESKTOP / "\u57fa\u4e8e\u865a\u5e7b\u5f15\u64ce\u7684\u52a8\u4f5c\u89d2\u8272\u626e\u6f14\u6e38\u620f\u8bbe\u8ba1\u4e0e\u5b9e\u73b0_v4_fixed.docx"
DST = DESKTOP / "\u57fa\u4e8e\u865a\u5e7b\u5f15\u64ce\u7684\u52a8\u4f5c\u89d2\u8272\u626e\u6f14\u6e38\u620f\u8bbe\u8ba1\u4e0e\u5b9e\u73b0_v5.docx"

# ── PlantUML encoding ──────────────────────────────────────────
_CHARS = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz-_"

def _encode6(data: bytes) -> str:
    out = []
    for i in range(0, len(data), 3):
        b = data[i:i+3]
        if len(b) == 1:   b += b'\x00\x00'
        elif len(b) == 2: b += b'\x00'
        n = (b[0] << 16) | (b[1] << 8) | b[2]
        out += [_CHARS[(n>>18)&63], _CHARS[(n>>12)&63],
                _CHARS[(n>> 6)&63], _CHARS[n&63]]
    return "".join(out)

def plantuml_url(src: str) -> str:
    compressed = zlib.compress(src.encode("utf-8"))[2:-4]
    return "https://www.plantuml.com/plantuml/png/" + _encode6(compressed)

def download_png(src: str, path: pathlib.Path) -> bool:
    url = plantuml_url(src)
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as r:
            data = r.read()
        if data[:4] == b'\x89PNG':
            path.write_bytes(data)
            print(f"  OK: {path.name} ({len(data)} bytes)")
            return True
        print(f"  FAIL: not PNG")
        return False
    except Exception as e:
        print(f"  ERROR: {e}")
        return False

# ── helpers ────────────────────────────────────────────────────
SZ_X4 = Pt(12)

def set_run_font(run, cn="\u5b8b\u4f53", en="Times New Roman", size=SZ_X4, bold=False):
    run.bold = bold
    run.font.size = size
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    fonts_el = rPr.find(qn('w:rFonts'))
    if fonts_el is None:
        fonts_el = OxmlElement('w:rFonts')
        rPr.insert(0, fonts_el)
    fonts_el.set(qn('w:eastAsia'), cn)
    fonts_el.set(qn('w:ascii'), en)
    fonts_el.set(qn('w:hAnsi'), en)

def set_cell_font(cell, text, bold=False, align='center'):
    """清空单元格内容，写入统一格式文本"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, bold=bold)

def set_table_borders(table):
    """设置表格全边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def make_caption(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '21')
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '21')
    rPr.append(sz); rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    r.append(t); p.append(r)
    return p

def make_image_para(img_path, width_cm=13):
    tmp_doc = Document()
    p = tmp_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Cm(width_cm))
    return p._element

# ── data dictionary tables ─────────────────────────────────────
TABLES = [
    # (title_keyword, rows)
    ("\u88684-5-1", [
        ("\u5b57\u6bb5\u540d", "\u6570\u636e\u7c7b\u578b", "\u8bf4\u660e"),
        ("CharacterID", "FString", "\u89d2\u8272\u552f\u4e00\u6807\u8bc6"),
        ("CharacterClass", "TSubclassOf<APawn>", "\u89d2\u8272Pawn\u7c7b"),
        ("CharacterName", "FText", "\u89d2\u8272\u540d\u79f0"),
        ("CharacterAttributes", "TArray<FCharacterAttributeModifier>", "\u89d2\u8272\u5c5e\u6027\u5217\u8868\uff08\u5305\u542bAttributeName\u3001AttributeValue\uff09"),
        ("CharacterPreviewAnimAsset", "TObjectPtr<UAnimMontage>", "\u9884\u89c8\u52a8\u753b\u8d44\u6e90"),
        ("CharacterInitData", "TSubclassOf<UGameplayEffect>", "\u521d\u59cb\u5316\u5c5e\u6027GameplayEffect"),
        ("CharacterBuffEffect", "TSubclassOf<UGameplayEffect>", "Buff\u6548\u679cGameplayEffect"),
        ("CharacterMainAbility", "TSubclassOf<UGameplayAbility>", "\u4e3b\u80fd\u529b\u7c7b"),
        ("CharacterHurtAbility", "TSubclassOf<UGameplayAbility>", "\u53d7\u4f24\u80fd\u529b\u7c7b"),
    ]),
    ("\u88684-5-2", [
        ("\u5b57\u6bb5\u540d", "\u6570\u636e\u7c7b\u578b", "\u8bf4\u660e"),
        ("AIName", "FString", "AI\u540d\u79f0"),
        ("AIPawnClass", "TSubclassOf<APawn>", "AI Pawn\u7c7b"),
        ("AIHurtAbility", "TSubclassOf<UGameplayAbility>", "\u53d7\u4f24\u80fd\u529b\u7c7b"),
        ("AIHealingAbility", "TSubclassOf<UGameplayAbility>", "\u6cbb\u7597\u80fd\u529b\u7c7b"),
        ("AIInitData", "TSubclassOf<UGameplayEffect>", "\u521d\u59cb\u5316\u6570\u636eGameplayEffect"),
        ("IDConfigDataAsset", "UInGamePlayerConfig*", "\u914d\u7f6e\u6570\u636e\u8d44\u4ea7"),
        ("AIControllerClass", "TSubclassOf<AInGameAIController>", "AI\u63a7\u5236\u5668\u7c7b"),
        ("AITree", "UBehaviorTree*", "\u884c\u4e3a\u6811\u8d44\u6e90"),
        ("AIBlackboard", "UBlackboardData*", "\u9ed1\u677f\u6570\u636e"),
        ("BlackboardKeys", "FAIBlackboardKeys", "\u9ed1\u677f\u952e\u503c\u914d\u7f6e"),
        ("PatrollingPath", "TSubclassOf<AActor>", "\u5de1\u903b\u8def\u5f84"),
        ("EquipmentActors", "TArray<TSubclassOf<AActor>>", "\u88c5\u5907Actor\u7c7b\u5217\u8868"),
    ]),
    ("\u88684-5-3", [
        ("\u5b57\u6bb5\u540d", "\u6570\u636e\u7c7b\u578b", "\u8bf4\u660e"),
        ("WeaponName", "FString", "\u6b66\u5668\u540d\u79f0"),
        ("WeaponType", "TEnumAsByte<EWeaponType>", "\u6b66\u5668\u7c7b\u578b\uff08\u5f92\u624b/\u5251/\u76fe\uff09"),
        ("EquipmentBaseAbilityClasses", "TArray<TSubclassOf<UGameplayAbility>>", "\u57fa\u7840\u80fd\u529b\u7c7b\u5217\u8868"),
        ("SpecialAbilityConfig", "TArray<FEquipmentSpecialAbilityConfig>", "\u7279\u6b8a\u80fd\u529b\u914d\u7f6e\uff08ActivateType+AbilityClass\uff09"),
        ("DetectSocketNames", "TArray<FName>", "\u78b0\u649e\u68c0\u6d4bSocket\u540d\u5217\u8868"),
        ("WeaponClass", "TSubclassOf<AActor>", "\u6b66\u5668Actor\u7c7b"),
        ("ComboWindow", "float", "\u8fde\u62db\u7a97\u53e3\u65f6\u95f4\uff08\u79d2\uff09"),
    ]),
    ("\u88684-5-4", [
        ("\u5b57\u6bb5\u540d", "\u6570\u636e\u7c7b\u578b", "\u8bf4\u660e"),
        ("MapID", "FString", "\u5730\u56fe\u552f\u4e00\u6807\u8bc6"),
        ("LevelToOpen", "TSoftObjectPtr<UWorld>", "\u5173\u5361\u8d44\u6e90\u5f15\u7528"),
        ("GameModeClassForLevel", "TSubclassOf<AGameModeBase>", "\u5173\u5361GameMode\u7c7b"),
        ("GameModeName", "FString", "GameMode\u540d\u79f0"),
        ("MapIcon", "TObjectPtr<UTexture2D>", "\u5730\u56fe\u56fe\u6807"),
        ("SpawnPoints", "TArray<FSpawnPointData>", "\u73a9\u5bb6\u51fa\u751f\u70b9\u6570\u7ec4\uff08\u4f4d\u7f6e/\u542f\u7528/\u5bb9\u91cf\uff09"),
        ("EscapePoints", "TArray<FEscapePointData>", "\u64a4\u79bb\u70b9\u6570\u7ec4\uff08\u7b49\u5f85\u65f6\u95f4/\u7269\u54c1\u9700\u6c42/Portal\u7c7b\uff09"),
    ]),
    ("\u88684-5-5", [
        ("\u5b57\u6bb5\u540d", "\u6570\u636e\u7c7b\u578b", "\u8bf4\u660e"),
        ("OperationAbilities", "TArray<TSubclassOf<UGameplayAbility>>", "\u64cd\u4f5c\u80fd\u529b\u7c7b\u5217\u8868"),
        ("OperationStaminaCost", "float", "\u64cd\u4f5c\u4f53\u529b\u6d88\u8017"),
        ("StaminaRecoverWindow", "float", "\u4f53\u529b\u6062\u590d\u7a97\u53e3\u65f6\u95f4\uff08\u79d2\uff09"),
    ]),
    ("\u88684-5-6", [
        ("\u5b57\u6bb5\u540d", "\u6570\u636e\u7c7b\u578b", "\u8bf4\u660e"),
        ("WidgetID", "FString", "UI\u552f\u4e00\u6807\u8bc6"),
        ("WidgetClass", "TSubclassOf<UUserWidget>", "Widget\u7c7b"),
        ("WidgetOrder", "int32", "\u5c42\u7ea7\u663e\u793a\u987a\u5e8f"),
    ]),
]

# ── PlantUML sources ───────────────────────────────────────────
DFD_SRC = """\
@startuml
skinparam defaultFontName Arial
skinparam defaultFontSize 12
skinparam rectangle {
  BackgroundColor<<DataSource>> #E3F2FD
  BackgroundColor<<Process>> #FFF3E0
  BackgroundColor<<Store>> #E8F5E9
  BorderColor #333333
}

rectangle "CSV/JSON\\n\u914d\u7f6e\u6587\u4ef6" as Config <<DataSource>>
rectangle "UPlayerGameInstance\\n\u5f02\u6b65\u52a0\u8f7d" as GI <<Process>>
rectangle "\u7f13\u5b58\u533a\\n(FCharacterData,\\nFMapManifest\u7b49)" as Cache <<Store>>

rectangle "\u89d2\u8272\u9009\u62e9" as CharSel
rectangle "\u5546\u5e97\u7cfb\u7edf" as Shop
rectangle "\u6b66\u5668\u7cfb\u7edf" as Weapon
rectangle "\u5173\u5361\u7cfb\u7edf" as Level
rectangle "AI\u7cfb\u7edf" as AI
rectangle "UI\u7cfb\u7edf" as UI

Config --> GI : \u8bfb\u53d6DataTable
GI --> Cache : \u89e3\u6790\u4e0e\u7f13\u5b58
Cache --> CharSel : \u89d2\u8272\u914d\u7f6e
Cache --> Shop : \u5546\u54c1\u914d\u7f6e
Cache --> Weapon : \u6b66\u5668\u914d\u7f6e
Cache --> Level : \u5730\u56fe\u914d\u7f6e
Cache --> AI : AI\u914d\u7f6e
Cache --> UI : UI\u914d\u7f6e
@enduml
"""

ER_SRC = """\
@startuml
skinparam class {
  BackgroundColor #FEFEFE
  BorderColor #333333
}

entity "\u73a9\u5bb6 Player" as Player {
  PlayerUserName : FString
  PlayerTotalMoney : float
  HealthMod : float
  AttackPowerMod : float
  DefensePowerMod : float
}

entity "\u89d2\u8272 Character" as Character {
  CharacterID : FString
  CharacterName : FText
  CharacterClass : TSubclassOf<APawn>
}

entity "\u6b66\u5668 Weapon" as Weapon {
  WeaponName : FString
  WeaponType : EWeaponType
  ComboWindow : float
}

entity "\u5730\u56fe Map" as Map {
  MapID : FString
  GameModeName : FString
}

entity "AI\u89d2\u8272 AICharacter" as AIChar {
  AIName : FString
  AIPawnClass : TSubclassOf<APawn>
}

entity "\u5546\u5e97\u7269\u54c1 ShopItem" as ShopItem {
  ItemID : FString
  ItemPrice : float
  AttackPowerUp : float
  DefensePowerUp : float
  HealthPowerUp : float
}

entity "UI\u914d\u7f6e UIWidget" as UIW {
  WidgetID : FString
  WidgetOrder : int32
}

entity "\u51fa\u751f\u70b9 SpawnPoint" as SP {
  PointTransform : FTransform
  MaxPlayerToSpawn : int32
}

entity "\u64a4\u79bb\u70b9 EscapePoint" as EP {
  EscapeWaitingTime : float
  bNeedEscapeItem : bool
}

Player ||--o{ Character : \u62e5\u6709
Character ||--o| Weapon : \u88c5\u5907
Map ||--o{ SP : \u5305\u542b
Map ||--o{ EP : \u5305\u542b
Map ||--o{ AIChar : \u751f\u6210
Player ||--o{ ShopItem : \u8d2d\u4e70
UIW ||--o{ Player : \u663e\u793a
@enduml
"""

# ── main ───────────────────────────────────────────────────────
UML_DIR.mkdir(parents=True, exist_ok=True)
dfd_path = UML_DIR / "4_5_dfd.png"
er_path  = UML_DIR / "4_5_er.png"

print("Step 1: Download diagrams...")
ok1 = download_png(DFD_SRC, dfd_path)
ok2 = download_png(ER_SRC, er_path)
if not (ok1 and ok2):
    print("Diagram download failed, aborting.")
    raise SystemExit(1)

print(f"\nStep 2: Open thesis...")
doc = Document(str(SRC))

# Build all tables first (they get appended at end)
# We will move them later via XML
print("Step 3: Build tables...")
table_elements = {}  # keyword -> (table_element, caption_element)
for keyword, rows in TABLES:
    tbl = doc.add_table(rows=len(rows), cols=3)
    set_table_borders(tbl)
    for ri, (c1, c2, c3) in enumerate(rows):
        set_cell_font(tbl.rows[ri].cells[0], c1, bold=(ri==0), align='left' if ri>0 else 'center')
        set_cell_font(tbl.rows[ri].cells[1], c2, bold=(ri==0), align='left' if ri>0 else 'center')
        set_cell_font(tbl.rows[ri].cells[2], c3, bold=(ri==0), align='left' if ri>0 else 'center')
    # caption
    cap = make_caption(f"{keyword}  \u6570\u636e\u5b57\u5178")
    table_elements[keyword] = (tbl._tbl, cap)

print("Step 4: Insert tables and images...")
# Find insertion points
para_map = {}
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    for keyword, _ in TABLES:
        if keyword in t:
            para_map[keyword] = para
            break
    if t == "\u88684-5-6 UI\u754c\u9762\u914d\u7f6e":
        para_map["last_table_title"] = para

# Insert tables after each title paragraph
for keyword, (tbl_elem, cap_elem) in table_elements.items():
    if keyword not in para_map:
        print(f"  WARN: not found {keyword}")
        continue
    para_elem = para_map[keyword]._element
    # insert table then caption after the title paragraph
    para_elem.addnext(tbl_elem)
    tbl_elem.addnext(cap_elem)
    print(f"  Inserted {keyword}")

# Insert DFD before first table (after 4.5.2 heading)
heading_452 = None
for para in doc.paragraphs:
    if "4.5.2" in para.text and "\u6570\u636e\u8868\u8bbe\u8ba1" in para.text:
        heading_452 = para
        break

if heading_452 and para_map:
    first_key = "\u88684-5-1"
    if first_key in para_map:
        first_para = para_map[first_key]
        img_elem = make_image_para(dfd_path, width_cm=12)
        cap_elem = make_caption("\u56fe4-5-1  \u6570\u636e\u8868\u52a0\u8f7d\u6570\u636e\u6d41\u56fe")
        first_para._element.addprevious(cap_elem)
        first_para._element.addprevious(img_elem)
        print("  Inserted DFD before first table")

# Insert ER after last table
if "last_table_title" in para_map:
    last_para = para_map["last_table_title"]
    # find the paragraph after the last table (which is the empty paragraph after table title)
    # We need to find the table that was inserted after last_table_title
    # and insert ER after that table
    # The table was inserted as XML sibling after last_para._element
    # Find next non-empty paragraph after last table
    found_last = False
    insert_after = None
    for para in doc.paragraphs:
        if para._element is last_para._element:
            found_last = True
            continue
        if found_last:
            # Skip empty paragraphs
            if para.text.strip():
                insert_after = para
                break
    if insert_after is None:
        # fallback: insert before chapter 5
        for para in doc.paragraphs:
            if "\u7b2c5\u7ae0" in para.text:
                insert_after = para
                break
    if insert_after:
        img_elem = make_image_para(er_path, width_cm=13)
        cap_elem = make_caption("\u56fe4-5-2  \u6e38\u620f\u914d\u7f6e\u6570\u636eE-R\u56fe")
        insert_after._element.addprevious(cap_elem)
        insert_after._element.addprevious(img_elem)
        print("  Inserted E-R diagram")

# Save
try:
    doc.save(str(DST))
    print(f"\nSaved: {DST.name}")
except PermissionError:
    alt = DST.with_stem(DST.stem + "_alt")
    doc.save(str(alt))
    print(f"\nPermissionError, saved as: {alt.name}")
