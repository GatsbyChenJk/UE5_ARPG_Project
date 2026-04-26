# -*- coding: utf-8 -*-
import zlib, urllib.request, pathlib
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DESKTOP = pathlib.Path.home() / "Desktop" / "\u4e13\u8bbe\u6bd5\u8bbe\u6587\u4ef6"
UML_DIR = DESKTOP / "uml"
V3 = DESKTOP / "\u57fa\u4e8e\u865a\u5e7b\u5f15\u64ce\u7684\u52a8\u4f5c\u89d2\u8272\u626e\u6f14\u6e38\u620f\u8bbe\u8ba1\u4e0e\u5b9e\u73b0_v3.docx"
V4 = DESKTOP / "\u57fa\u4e8e\u865a\u5e7b\u5f15\u64ce\u7684\u52a8\u4f5c\u89d2\u8272\u626e\u6f14\u6e38\u620f\u8bbe\u8ba1\u4e0e\u5b9e\u73b0_v4.docx"

# ── PlantUML encoding ──────────────────────────────────────────
_CHARS = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz-_"

def _encode6(data: bytes) -> str:
    out = []
    for i in range(0, len(data), 3):
        b = data[i:i+3]
        if len(b) == 1:   b += b'\x00\x00'
        elif len(b) == 2: b += b'\x00'
        n = (b[0] << 16) | (b[1] << 8) | b[2]
        out += [_CHARS[(n>>18)&63], _CHARS[(n>>12)&63],
                _CHARS[(n>> 6)&63], _CHARS[n&63]]
    return "".join(out)

def plantuml_url(src: str) -> str:
    compressed = zlib.compress(src.encode("utf-8"))[2:-4]
    return "https://www.plantuml.com/plantuml/png/" + _encode6(compressed)

def download_png(src: str, path: pathlib.Path) -> bool:
    url = plantuml_url(src)
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as r:
            data = r.read()
        if data[:4] == b'\x89PNG':
            path.write_bytes(data)
            print(f"  OK: {path.name}  ({len(data)} bytes)")
            return True
        print(f"  FAIL (not PNG): {data[:200]}")
        return False
    except Exception as e:
        print(f"  ERROR: {e}")
        return False

# ── font helper ────────────────────────────────────────────────
SZ_X4 = Pt(12)

def make_caption(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz   = OxmlElement('w:sz');   sz.set(qn('w:val'),   '21')
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '21')
    b = OxmlElement('w:b')
    rPr.append(b); rPr.append(sz); rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    r.append(t); p.append(r)
    return p

def make_image_para(img_path, width_cm=13):
    tmp_doc = Document()
    p = tmp_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Cm(width_cm))
    return p._element

# ── PlantUML source for 3.3 use case diagram ──────────────────
USECASE_SRC = """\
@startuml
skinparam defaultFontName Arial
skinparam defaultFontSize 12
skinparam usecase {
  BackgroundColor #FEFEFE
  BorderColor #333333
}
skinparam actor {
  BorderColor #333333
}
skinparam ArrowColor #555555

actor "\u73a9\u5bb6" as Player

rectangle "\u57fa\u4e8e\u865a\u5e7b\u5f15\u64ce\u7684ARPG\u6e38\u620f\u7cfb\u7edf" {
  package "\u5c40\u5916\u5927\u5385" {
    usecase (\u767b\u5f55\u5927\u5385) as UC1
    usecase (\u9009\u62e9\u89d2\u8272) as UC2
    usecase (\u5546\u5e97\u8d2d\u7269) as UC8
    usecase (\u67e5\u770b\u5b58\u6863) as UC9
  }
  package "\u5c40\u5185\u6e38\u620f" {
    usecase (\u8fdb\u5165\u5173\u5361) as UC3
    usecase (\u63a2\u7d22\u573a\u666f) as UC4
    usecase (\u53d1\u8d77\u6218\u6597) as UC5
    usecase (\u62fe\u53d6\u7269\u54c1) as UC6
    usecase (\u64a4\u79bb\u5173\u5361) as UC7
  }
}

Player --> UC1
Player --> UC2
Player --> UC3
Player --> UC4
Player --> UC5
Player --> UC6
Player --> UC7
Player --> UC8
Player --> UC9

UC2 .up.> UC1 : <<include>>
UC8 .up.> UC1 : <<include>>
UC3 .up.> UC2 : <<include>>
UC5 .> UC4 : <<extend>>
UC6 .> UC5 : <<extend>>
UC7 .> UC4 : <<include>>
@enduml
"""

# ── main ───────────────────────────────────────────────────────
UML_DIR.mkdir(parents=True, exist_ok=True)
png_path = UML_DIR / "3_3_usecase.png"

print("Downloading use case diagram...")
if not download_png(USECASE_SRC, png_path):
    raise SystemExit(1)

print(f"\nOpening {V3.name} ...")
doc = Document(str(V3))

# Find 3.3 section heading
heading_33 = None
for para in doc.paragraphs:
    if "3.3" in para.text and "\u7cfb\u7edf\u7528\u4f8b" in para.text:
        heading_33 = para
        break

if heading_33 is None:
    print("ERROR: Cannot find 3.3 heading")
    raise SystemExit(1)

# Find insertion point: page break or 第4章 heading after 3.3
insert_before = None
past_33 = False
for para in doc.paragraphs:
    if para._element is heading_33._element:
        past_33 = True
        continue
    if not past_33:
        continue
    xml = para._element.xml
    # page break or next chapter heading
    if 'w:br' in xml or "\u7b2c4\u7ae0" in para.text:
        insert_before = para
        break

if insert_before is None:
    print("ERROR: Cannot find insertion point after 3.3")
    raise SystemExit(1)

print(f"Inserting before: [{insert_before.text[:40]}]")

img_elem = make_image_para(png_path, width_cm=13)
cap_elem = make_caption("\u56fe3-1  \u7cfb\u7edf\u7528\u4f8b\u56fe")

insert_before._element.addprevious(img_elem)
insert_before._element.addprevious(cap_elem)

try:
    doc.save(str(V4))
    print(f"\nSaved: {V4.name}")
except PermissionError:
    V4b = V4.with_stem(V4.stem + "b")
    doc.save(str(V4b))
    print(f"\nPermissionError on v4, saved as: {V4b.name}")
    print("Please close _v4.docx in Word, then rename _v4b.docx to _v4.docx if needed.")
