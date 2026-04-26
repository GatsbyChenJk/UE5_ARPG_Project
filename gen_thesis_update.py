# -*- coding: utf-8 -*-
"""
向已生成的论文 _v2.docx 补充：
  5.3.4  伤害系统实现（含伤害时序图 图5-11）
  5.3.5  局内物品系统实现（含物品时序图 图5-12）
  在 5.2.1 末尾插入角色选择时序图（图5-9）
  在 5.2.2 末尾插入商店时序图（图5-10）
输出 _v3.docx，并刷新附件 docx
"""

import zlib, copy, time, sys, urllib.request
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────
# 路径
# ──────────────────────────────────────────
DESKTOP  = Path(r"C:\Users\25768\Desktop\专设毕设文件")
V2 = DESKTOP / "\u57fa\u4e8e\u865a\u5e7b\u5f15\u64ce\u7684\u52a8\u4f5c\u89d2\u8272\u626e\u6f14\u6e38\u620f\u8bbe\u8ba1\u4e0e\u5b9e\u73b0_v2.docx"
V3 = DESKTOP / "\u57fa\u4e8e\u865a\u5e7b\u5f15\u64ce\u7684\u52a8\u4f5c\u89d2\u8272\u626e\u6f14\u6e38\u620f\u8bbe\u8ba1\u4e0e\u5b9e\u73b0_v3.docx"
APP_PATH = DESKTOP / "\u57fa\u4e8e\u865a\u5e7b\u5f15\u64ce\u7684\u52a8\u4f5c\u89d2\u8272\u626e\u6f14\u6e38\u620f\u8bbe\u8ba1\u4e0e\u5b9e\u73b0-\u9644\u4ef6.docx"
UML_DIR  = DESKTOP / "uml"
UML_DIR.mkdir(parents=True, exist_ok=True)

SZ_3  = Pt(16)
SZ_4  = Pt(14)
SZ_X4 = Pt(12)
SZ_5  = Pt(10.5)

if not V2.exists():
    ORIG = DESKTOP / "\u57fa\u4e8e\u865a\u5e7b\u5f15\u64ce\u7684\u52a8\u4f5c\u89d2\u8272\u626e\u6f14\u6e38\u620f\u8bbe\u8ba1\u4e0e\u5b9e\u73b0.docx"
    if ORIG.exists():
        V2 = ORIG
    else:
        print("Error: cannot find thesis docx")
        sys.exit(1)

# ──────────────────────────────────────────
# PlantUML
# ──────────────────────────────────────────
def _e6(b):
    if b < 10:  return chr(48 + b)
    b -= 10
    if b < 26:  return chr(65 + b)
    b -= 26
    if b < 26:  return chr(97 + b)
    b -= 26
    return '-' if b == 0 else '_'

def _a3(b1, b2, b3):
    return _e6(b1>>2) + _e6(((b1&3)<<4)|(b2>>4)) + _e6(((b2&0xF)<<2)|(b3>>6)) + _e6(b3&0x3F)

def encode_puml(text):
    d = zlib.compress(text.encode('utf-8'), 9)[2:-4]
    out, i = '', 0
    while i + 2 < len(d): out += _a3(d[i], d[i+1], d[i+2]); i += 3
    r = len(d) - i
    if r == 1:   out += _a3(d[i], 0, 0)[:2]
    elif r == 2: out += _a3(d[i], d[i+1], 0)[:3]
    return out

def download(puml, path, retries=4):
    url = 'https://www.plantuml.com/plantuml/png/' + encode_puml(puml)
    for n in range(retries):
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=30) as r:
                path.write_bytes(r.read())
            print(f'  OK: {path.name}')
            return True
        except Exception as e:
            print(f'  attempt {n+1} failed: {e}')
            time.sleep(3)
    print(f'  FAILED: {path.name}')
    return False

# ──────────────────────────────────────────
# 4 张时序图源码
# ──────────────────────────────────────────
SEQ_SKINS = """skinparam sequenceArrowThickness 1.5
skinparam sequenceParticipantBackgroundColor #F8F9FA
skinparam sequenceParticipantBorderColor #495057
skinparam sequenceActorBackgroundColor #E8F4FD
skinparam sequenceLifeLineBorderColor #888888
skinparam sequenceParticipantFontSize 10
"""

SEQ_DIAGRAMS = [
    ("5_2_1_sequence",
     "\u56fe5-9\u2002\u89d2\u8272\u9009\u62e9\u6a21\u5757\u65f6\u5e8f\u56fe",
     "\u89d2\u8272\u9009\u62e9\u5b9e\u73b0",
     f"""@startuml
{SEQ_SKINS}
title Character Select Sequence
actor "Player" as P
participant "UCharacterSelect" as CS
participant "ARPGSaveGameManager" as SGM
database "CharacterTable" as CT
participant "AInGameCharacter" as Char

P -> CS: Open lobby
CS -> SGM: GetTotalMoney()
SGM --> CS: PlayerTotalMoney
CS -> CT: QueryAllRows()
CT --> CS: TArray<FCharacterDataRow>
CS -> CS: Build character cards
P -> CS: Click character card
CS -> CS: selectedRowHandle = RowHandle
P -> CS: Enter game
CS -> SGM: SaveSelectedCharacter()
SGM -> SGM: AsyncSave(SlotName)

== Enter in-game level ==

Char -> CT: QueryRow(selectedRowHandle)
CT --> Char: FCharacterDataRow
Char -> Char: Init GAS AttributeSet\\n(Health/Attack/Defense)
@enduml"""),

    ("5_2_2_sequence",
     "\u56fe5-10\u2002\u5546\u5e97\u6a21\u5757\u65f6\u5e8f\u56fe",
     "\u5546\u5e97\u5b9e\u73b0",
     f"""@startuml
{SEQ_SKINS}
title Shop Module Sequence
actor "Player" as P
participant "UARPGLeaf_ShopItem" as Shop
participant "ARPGSaveGameManager" as SGM
participant "ARPGEventManager" as EM
database "ShopItemDataTable" as ST

P -> Shop: Open shop
Shop -> ST: QueryAllRows()
ST --> Shop: TArray<FShopItemDataRow>
Shop -> SGM: GetTotalMoney()
SGM --> Shop: PlayerTotalMoney
Shop -> Shop: Render item cards

P -> Shop: Click Buy button
Shop -> Shop: Check Price <= Balance
alt Balance sufficient
  Shop -> SGM: DeductMoney(Price) [Server RPC]
  SGM -> SGM: PlayerTotalMoney -= Price
  SGM -> SGM: AsyncSave(SlotName)
  SGM --> Shop: Success
  Shop -> EM: Broadcast(ARPG_MoneyUpdated)
  EM --> Shop: Notify balance display
else Balance insufficient
  Shop -> P: Show insufficient funds hint
end
@enduml"""),

    ("5_3_4_sequence",
     "\u56fe5-11\u2002\u4f24\u5bb3\u7cfb\u7edf\u65f6\u5e8f\u56fe",
     "\u4f24\u5bb3\u7cfb\u7edf\u5b9e\u73b0",
     f"""@startuml
{SEQ_SKINS}
title Damage System Sequence
actor "Player" as P
participant "AInGameCharacter" as Char
participant "UCharacterAttackComponent" as AC
participant "UCharacterDamageComponent" as DC
participant "AbilitySystemComponent" as ASC
participant "UUExec_DamageCalculation" as EC
participant "UInGameCharacterAttributeSet" as Attr
participant "AInGameAICharacter" as AI

P -> Char: Press attack key
Char -> AC: OnCharacterAttack()
AC -> AC: Server_CharacterHandleAttack()\\ncombo++
AC -> AC: IsValidStamina()
AC -> ASC: TryActivateAbility(AttackAbility)
ASC -> DC: Trigger overlap detection
DC -> DC: SphereOverlap OR CapsuleOverlap\\n(ECC_GameTraceChannel3)
DC -> DC: UpdateHitResult() -> StoredHitActors
DC -> DC: ApplyDamageToHitActors()
DC -> ASC: SendGameplayEventToActor(\\n  "CharacterAttribute.Health.Consume",\\n  Magnitude=AttackPower)
ASC -> EC: Execute_Implementation()
EC -> EC: FinalDmg = AttackPower\\n - ArmorDefense - DefensePower\\n = Max(0, FinalDmg)
EC -> Attr: AddOutputModifier(Health, -FinalDmg)
Attr -> Attr: Health updated & replicated

alt Health <= 0
  Attr -> AI: CharacterDeath()
  AI -> AI: Destroy weapons -> Destroy self
else Health > 0
  Attr -> Char: OnRep_Health() -> Refresh HUD
end
@enduml"""),

    ("5_3_5_sequence",
     "\u56fe5-12\u2002\u5c40\u5185\u7269\u54c1\u7cfb\u7edf\u65f6\u5e8f\u56fe",
     "\u5c40\u5185\u7269\u54c1\u7cfb\u7edf\u5b9e\u73b0",
     f"""@startuml
{SEQ_SKINS}
title Inventory System Sequence
actor "Player" as P
participant "AInventoryPlayerController" as PC
participant "UItemComponent" as IC
participant "UInventoryComp" as Inv
participant "FInventoryFastArray" as FA
participant "AbilitySystemComponent" as ASC
participant "ARPGEventManager" as EM

group Item Pickup
  P -> PC: Press E (interact)
  PC -> IC: PrimaryInteract()
  IC -> IC: PickUp() -> Destroy(Owner)
  PC -> Inv: TryAddItem(ItemComp)
  Inv -> Inv: HasRoomForItem()
  Inv -> FA: ServerAddNewItem() [Server RPC]
  FA -> FA: AddEntry() -> MarkItemDirty()
  FA --> Inv: PostReplicatedAdd()\\n-> OnItemAdded.Broadcast()
  Inv --> P: Grid UI refreshes
end

group Item Consume
  P -> Inv: Server_ConsumeItem(Item) [RPC]
  Inv -> Inv: StackCount--\\nif 0: RemoveEntry()
  Inv -> IC: FConsumableFragment::OnConsume()
  IC -> ASC: ApplyGameplayEffectByValue(\\n  ModifyEffectClass,\\n  SetByCallerMagnitude=Value)
  ASC -> ASC: Update Health/Mana attribute
end

group Item Equip
  P -> Inv: Server_EquipSlotClicked(Item) [RPC]
  Inv -> Inv: Multicast_EquipSlotClicked()
  Inv -> EM: ARPG_EVENT "CharacterEquipWeapon"\\n(UARPGEventData_Weapon)
  EM -> Char: Update attack config & anim
end
@enduml"""),
]

# ──────────────────────────────────────────
# 下载时序图
# ──────────────────────────────────────────
print("=" * 55)
print("Step 1: Download sequence diagram PNGs")
print("=" * 55)
downloaded = {}
for prefix, caption, _, puml in SEQ_DIAGRAMS:
    out = UML_DIR / f"{prefix}.png"
    print(f"\n{caption}")
    if download(puml, out):
        downloaded[prefix] = out

# ──────────────────────────────────────────
# 段落构建辅助
# ──────────────────────────────────────────

def ensure_rpr(run):
    rPr = run._element.get_or_add_rPr()
    fonts = rPr.find(qn('w:rFonts'))
    if fonts is None:
        fonts = OxmlElement('w:rFonts')
        rPr.insert(0, fonts)
    return rPr, fonts

def set_run_font(run, cn='Songti SC', en='Times New Roman', size=SZ_X4, bold=False):
    run.bold  = bold
    run.font.size = size
    run.font.name = en
    _, f = ensure_rpr(run)
    f.set(qn('w:eastAsia'), cn)

def make_heading2(doc_ctx, text):
    """Create a heading-2 paragraph element (not attached to doc)"""
    from docx import Document as D2
    tmp = D2()
    p = tmp.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    set_run_font(r, cn='\u9ed1\u4f53', en='Times New Roman', size=SZ_X4, bold=True)
    return copy.deepcopy(p._element)

def make_heading3(doc_ctx, text):
    from docx import Document as D2
    tmp = D2()
    p = tmp.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    set_run_font(r, cn='\u9ed1\u4f53', en='Times New Roman', size=SZ_X4, bold=False)
    return copy.deepcopy(p._element)

def make_body(text, first_indent=True):
    from docx import Document as D2
    tmp = D2()
    p = tmp.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_indent:
        pf.first_line_indent = Pt(24)
    r = p.add_run(text)
    set_run_font(r, cn='\u5b8b\u4f53', en='Times New Roman', size=SZ_X4)
    return copy.deepcopy(p._element)

def make_code(text):
    from docx import Document as D2
    tmp = D2()
    p = tmp.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.left_indent  = Pt(24)
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = SZ_5
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    p._p.pPr.append(shd)
    return copy.deepcopy(p._element)

def make_image(img_path, width_inches=4.6):
    from docx import Document as D2
    tmp = D2()
    p = tmp.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)
    p.add_run().add_picture(str(img_path), width=Inches(width_inches))
    return copy.deepcopy(p._element)

def make_caption(text):
    p_xml = OxmlElement('w:p')
    pPr   = OxmlElement('w:pPr')
    jc    = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
    sp    = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '120')
    pPr.append(jc); pPr.append(sp); p_xml.append(pPr)
    r  = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz  = OxmlElement('w:sz');   sz.set(qn('w:val'),   '21')
    szC = OxmlElement('w:szCs'); szC.set(qn('w:val'),  '21')
    fnt = OxmlElement('w:rFonts')
    fnt.set(qn('w:ascii'),    'Times New Roman')
    fnt.set(qn('w:eastAsia'), '\u5b8b\u4f53')
    rPr.append(fnt); rPr.append(sz); rPr.append(szC)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    r.append(t); p_xml.append(r)
    return p_xml

def make_placeholder(text):
    p_xml = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    col = OxmlElement('w:color'); col.set(qn('w:val'), 'FF0000')
    rPr.append(col); r.append(rPr)
    t = OxmlElement('w:t'); t.text = f'[image not downloaded: {text}]'
    r.append(t); p_xml.append(r)
    return p_xml

def insert_image_block(ref_elem, prefix, caption):
    """Insert image + caption BEFORE ref_elem"""
    cap_elem = make_caption(caption)
    img_path = UML_DIR / f"{prefix}.png"
    if img_path.exists() and prefix in downloaded:
        img_elem = make_image(img_path)
    else:
        img_elem = make_placeholder(f"{prefix}.png")
    ref_elem.addprevious(cap_elem)
    ref_elem.addprevious(img_elem)
    # Reorder so image comes before caption
    cap_elem.addprevious(img_elem)

def insert_elements_before(ref_elem, elements):
    """Insert list of elements in order before ref_elem"""
    for elem in elements:
        ref_elem.addprevious(elem)
        ref_elem = elem   # move cursor so next element goes after this one
        # Wait, this is wrong. Let me rethink.
    # Actually: each addprevious inserts before ref_elem, but ref_elem stays the same.
    # Result order: elements[0], elements[1], ..., elements[-1], ref_elem  ✓

def insert_list_before(ref_elem, elements):
    """Insert elements in document order before ref_elem"""
    for elem in elements:
        ref_elem.addprevious(elem)
    # This inserts: elements[0] ... elements[-1] ref_elem? No!
    # addprevious always inserts immediately before ref_elem.
    # After e0: e0, ref
    # After e1 (addprevious(ref)): e0, e1, ref   ← WRONG, puts e1 between e0 and ref
    # Actually: each time we call ref.addprevious(eN), eN goes right before ref.
    # So: init: ref
    # add e0: e0, ref
    # add e1: e0, e1, ref  <- e1 is inserted right before ref, after e0
    # add e2: e0, e1, e2, ref
    # This IS the correct order! The key insight: addprevious moves eN to be
    # the immediate previous sibling of ref each time, so they accumulate in order.

# ──────────────────────────────────────────
# Section 5.3.4 content elements
# ──────────────────────────────────────────

def build_534():
    """Return list of XML elements for 5.3.4 section"""
    elems = []

    elems.append(make_heading2(None, "5.3.4\u3000\u4f24\u5bb3\u7cfb\u7edf\u5b9e\u73b0"))

    elems.append(make_body(
        "\u4f24\u5bb3\u7cfb\u7edf\u4ee5UCharacterDamageComponent\u4e3a\u6838\u5fc3\u7ec4\u4ef6\uff0c"
        "\u5c06\u78b0\u649e\u68c0\u6d4b\u3001\u4f24\u5bb3\u4e8b\u4ef6\u6d3e\u53d1\u4e0eGAS\u6267\u884c\u8ba1\u7b97\u4e09\u4e2a\u73af\u8282\u89e3\u8026\u5b9e\u73b0\u3002"
        "\u73a9\u5bb6\u6216AI\u89d2\u8272\u53d1\u52a8\u653b\u51fb\u65f6\uff0c\u7531\u7ec4\u4ef6\u8d1f\u8d23\u68c0\u6d4b\u547d\u4e2d\u76ee\u6807\u5e76\u751f\u6210GameplayEvent\uff0c"
        "GAS\u6846\u67b6\u63a5\u6536\u4e8b\u4ef6\u540e\u7531UUExec_DamageCalculation\u5b8c\u6210\u7cbe\u786e\u7684\u4f24\u5bb3\u6570\u503c\u8ba1\u7b97\uff0c"
        "\u6700\u7ec8\u901a\u8fc7\u5c5e\u6027\u96c6\u4fee\u6539\u53d7\u4f24\u65b9\u7684Health\u503c\u3002"
    ))

    elems.append(make_heading3(None, "5.3.4.1\u3000\u78b0\u649e\u68c0\u6d4b\u8bbe\u8ba1"))
    elems.append(make_body(
        "UCharacterDamageComponent\u5185\u90e8\u7ef4\u62a4StoredHitActors\u96c6\u5408\uff0c\u8bb0\u5f55\u5f53\u524d\u5e27\u7684\u547d\u4e2d\u76ee\u6807\u3002"
        "\u6839\u636e\u89d2\u8272\u662f\u5426\u6301\u6709\u6b66\u5668\uff0c\u5206\u522b\u91c7\u7528\u7403\u5f62\u68c0\u6d4b\uff08SphereOverlapDetection\uff09\u4e0e"
        "\u80f6\u56ca\u4f53\u68c0\u6d4b\uff08CapsuleOverlapDetection\uff09\u4e24\u79cd\u6a21\u5f0f\uff1a"
        "\u88f8\u624b\u653b\u51fb\u4ee5\u653b\u51fbSocket\u4e3a\u7403\u5fc3\u8fdb\u884c\u7403\u5f62\u68c0\u6d4b\uff1b"
        "\u6301\u6b66\u5668\u653b\u51fb\u5219\u4f9d\u636e\u6b66\u5668Socket\u7684\u7a7a\u95f4\u53d8\u6362\u8fdb\u884c\u80f6\u56ca\u4f53\u68c0\u6d4b\uff0c"
        "\u80fd\u66f4\u7cbe\u786e\u5730\u5339\u914d\u6b66\u5668\u653b\u51fb\u8303\u56f4\u3002"
        "\u4e24\u79cd\u6a21\u5f0f\u5747\u4f7f\u7528ECC_GameTraceChannel3\uff08\u8fd1\u6218\u4e13\u5c5e\u78b0\u649e\u901a\u9053\uff09\uff0c"
        "\u907f\u514d\u8bef\u89e6\u975e\u6218\u6597\u7269\u4f53\u3002\u68c0\u6d4b\u7ed3\u679c\u901a\u8fc7UpdateHitResult()\u53bb\u91cd\u540e\u5b58\u5165StoredHitActors\u3002"
    ))

    elems.append(make_heading3(None, "5.3.4.2\u3000\u4f24\u5bb3\u8ba1\u7b97\u4e0eGAS\u4e8b\u4ef6\u94fe"))
    elems.append(make_body(
        "ApplyDamageToHitActors()\u904d\u5386StoredHitActors\uff0c\u5c06\u653b\u51fb\u65b9\u5f53\u524d"
        "AttackPower\u5c5e\u6027\u503c\u5c01\u88c5\u81f3FGameplayEventData\uff0c"
        "\u8c03\u7528SendGameplayEventToActor\u5411\u53d7\u4f24\u65b9\u7684ASC\u53d1\u9001"
        "\u201cCharacterAttribute.Health.Consume\u201d\u4e8b\u4ef6\u3002\u53d7\u4f24\u65b9\u7684HurtAbility"
        "\u54cd\u5e94\u4e8b\u4ef6\u540e\u6fc0\u6d3bExecutionCalculation\uff0c"
        "\u7531UUExec_DamageCalculation::Execute_Implementation()\u6267\u884c\u4f24\u5bb3\u516c\u5f0f\uff1a"
    ))
    for line in [
        "  FinalDamage = TotalAttackPower - (EquipArmorDefense + SourceDefensePower)",
        "  FinalDamage = FMath::Max(0.0f, FinalDamage)",
    ]:
        elems.append(make_code(line))

    elems.append(make_body(
        "\u8ba1\u7b97\u7ed3\u679c\u901a\u8fc7AddOutputModifier\u4ee5\u8d1f\u503c\u5f62\u5f0f\u5199\u5165Health\u5c5e\u6027\uff0c"
        "\u89e6\u53d1\u7f51\u7edc\u590d\u5236\u4e0eUI\u5237\u65b0\u3002\u4e0b\u5217\u4e3a\u8be5\u8ba1\u7b97\u903b\u8f91\u7684\u6838\u5fc3\u5b9e\u73b0\uff1a"
    ))
    code_534 = [
        "void UUExec_DamageCalculation::Execute_Implementation(",
        "    const FGameplayEffectCustomExecutionParameters& ExecutionParams,",
        "    FGameplayEffectCustomExecutionOutput& OutExecutionOutput) const",
        "{",
        "    float TotalAttackPower = Spec.GetSetByCallerMagnitude(",
        "        FGameplayTag::RequestGameplayTag(",
        "            FName(\"Combat.WeaponAttack.BaseAttackPower\")), false, 0.f);",
        "    float EquipArmorDefense = Spec.GetSetByCallerMagnitude(",
        "        FGameplayTag::RequestGameplayTag(",
        "            FName(\"Combat.Defense.Armor\")), false, 0.f);",
        "    float SourceDefensePower = 0.f;",
        "    ExecutionParams.AttemptCalculateCapturedAttributeMagnitude(",
        "        AttributeStatics().DefensePowerDef,",
        "        FAggregatorEvaluateParameters(), SourceDefensePower);",
        "    float FinalDamage = FMath::Max(0.f,",
        "        TotalAttackPower - EquipArmorDefense - SourceDefensePower);",
        "    OutExecutionOutput.AddOutputModifier(",
        "        FGameplayModifierEvaluatedData(AttributeStatics().HealthProperty,",
        "            EGameplayModOp::Additive, -FinalDamage));",
        "}",
    ]
    for line in code_534:
        elems.append(make_code(line))

    elems.append(make_heading3(None, "5.3.4.3\u3000\u9632\u5fa1\u673a\u5236\u4e0e\u6b7b\u4ea1\u5904\u7406"))
    elems.append(make_body(
        "\u5f53\u89d2\u8272\u6fc0\u6d3b\u9632\u5fa1\u72b6\u6001\u65f6\uff08\u6301\u6709GameplayTag\u201cCombat.Defense.Activate\u201d\uff09\uff0c"
        "\u9632\u5fa1\u65b9\u7684DefensePower\u5c5e\u6027\u88ab\u7eb3\u5165\u4f24\u5bb3\u516c\u5f0f\u62b5\u6d88\u90e8\u5206\u653b\u51fb\u529b\uff0c"
        "\u540c\u65f6\u989d\u5916\u6d88\u8017DefenseStaminaCost\u4f53\u529b\uff0c\u6784\u6210\u653b\u5b88\u535a\u5f08\u7684\u6838\u5fc3\u673a\u5236\u3002"
        "DamageReduction\u5c5e\u6027\u4f5c\u4e3a\u9884\u7559\u6269\u5c55\u70b9\uff0c\u4f9b\u540e\u7eed\u62a4\u5177\u7cfb\u7edf\u6216\u88ab\u52a8\u6280\u80fd"
        "\u8fdb\u4e00\u6b65\u964d\u4f4e\u4f24\u5bb3\u3002"
        "Health\u5f52\u96f6\u65f6\u8c03\u7528CharacterDeath()\u865a\u51fd\u6570\uff1a"
        "AI\u89d2\u8272\u5148\u9500\u6bc1\u88c5\u5907\u7684\u6b66\u5668Actor\u518d\u9500\u6bc1\u81ea\u8eab\uff1b"
        "\u73a9\u5bb6\u89d2\u8272\u5219\u89e6\u53d1\u590d\u6d3b\u903b\u8f91\u3002"
    ))

    # Sequence diagram
    elems.append(make_body(
        "\u4f24\u5bb3\u7cfb\u7edf\u5404\u6a21\u5757\u7684\u5b8c\u6574\u8c03\u7528\u65f6\u5e8f\u5982\u56fe5-11\u6240\u793a\u3002"
    ))
    prefix_534 = "5_3_4_sequence"
    cap_534    = "\u56fe5-11\u2002\u4f24\u5bb3\u7cfb\u7edf\u65f6\u5e8f\u56fe"
    img_path_534 = UML_DIR / f"{prefix_534}.png"
    if img_path_534.exists() and prefix_534 in downloaded:
        elems.append(make_image(img_path_534))
    else:
        elems.append(make_placeholder(f"{prefix_534}.png"))
    elems.append(make_caption(cap_534))

    return elems

# ──────────────────────────────────────────
# Section 5.3.5 content elements
# ──────────────────────────────────────────

def build_535():
    elems = []
    elems.append(make_heading2(None, "5.3.5\u3000\u5c40\u5185\u7269\u54c1\u7cfb\u7edf\u5b9e\u73b0"))

    elems.append(make_body(
        "\u5c40\u5185\u7269\u54c1\u7cfb\u7edf\u4ee5\u63d2\u4ef6\u5f62\u5f0f\u5c01\u88c5\u5728Plugins/Inventory\u76ee\u5f55\u4e2d\uff0c"
        "\u901a\u8fc7Fragment\u7ec4\u5408\u6a21\u5f0f\u3001FFastArraySerializer\u7f51\u7edc\u540c\u6b65\u4e0eGAS\u96c6\u6210\uff0c"
        "\u5b9e\u73b0\u4e86\u7269\u54c1\u62fe\u53d6\u3001\u6d88\u8017\u3001\u88c5\u5907\u4e0e\u7a7a\u95f4\u80cc\u5305\u7ba1\u7406\u7684\u5b8c\u6574\u95ed\u73af\u3002"
    ))

    elems.append(make_heading3(None, "5.3.5.1\u3000Fragment\u7ec4\u5408\u6a21\u5f0f\u4e0e\u6570\u636e\u7ed3\u6784"))
    elems.append(make_body(
        "FItemManifest\u4ee5TInstancedStruct<FInventoryFragment>\u6570\u7ec4\u5b58\u50a8\u591a\u6001Fragment\uff0c"
        "\u6bcf\u4e2aFragment\u4ee3\u8868\u7269\u54c1\u7684\u4e00\u79cd\u7279\u6027\uff1a"
        "FImageFragment\u63d0\u4f9b\u56fe\u6807\u6570\u636e\uff0cFStackableFragment\u7ba1\u7406\u5806\u53e0\u6570\u91cf\uff0c"
        "FConsumableFragment\u5c01\u88c5\u6d88\u8017\u884c\u4e3a\uff0cFEquipmentFragment\u5b9a\u4e49\u88c5\u5907\u5c5e\u6027\u3002"
        "\u901a\u8fc7GetFragmentType<T>()\u6a21\u677f\u65b9\u6cd5\u53ef\u7c7b\u578b\u5b89\u5168\u5730\u67e5\u8be2\u4efb\u610fFragment\u3002"
        "\u8be5\u8bbe\u8ba1\u4f7f\u7269\u54c1\u529f\u80fd\u7684\u7ec4\u5408\u9ad8\u5ea6\u7075\u6d3b\uff0c"
        "\u65b0\u589e\u7269\u54c1\u7c7b\u578b\u53ea\u9700\u7ec4\u5408\u5df2\u6709Fragment\uff0c\u65e0\u9700\u4fee\u6539\u57fa\u7840\u7c7b\u4ee3\u7801\u3002"
    ))

    elems.append(make_heading3(None, "5.3.5.2\u3000\u80cc\u5305\u7ec4\u4ef6\u4e0e\u7f51\u7edc\u540c\u6b65"))
    elems.append(make_body(
        "UInventoryComp\u4f5c\u4e3aActorComponent\u9644\u52a0\u4e8e\u73a9\u5bb6\u89d2\u8272\uff0c\u662f\u7269\u54c1\u7cfb\u7edf\u7684\u670d\u52a1\u7aef\u6743\u5a01\u8282\u70b9\u3002"
        "\u5185\u90e8\u91c7\u7528FInventoryFastArray\uff08\u7ee7\u627f\u81eaFFastArraySerializer\uff09\u5b58\u50a8\u7269\u54c1\u5217\u8868\u3002"
        "FFastArraySerializer\u5b9e\u73b0\u4e86\u81ea\u52a8delta\u5e8f\u5217\u5316\uff0c"
        "\u4ec5\u5728\u7f51\u7edc\u66f4\u65b0\u65f6\u4f20\u8f93\u65b0\u589e/\u5220\u9664\u7684\u6761\u76ee\u5dee\u91cf\uff0c\u5927\u5e45\u964d\u4f4e\u80cc\u5305\u64cd\u4f5c\u7684\u7f51\u7edc\u6d41\u91cf\u3002"
        "PostReplicatedAdd\u4e0ePreReplicatedRemove\u9489\u5b50\u5728\u63a5\u6536\u7aef\u81ea\u52a8\u89e6\u53d1"
        "OnItemAdded/OnItemRemoved\u59d4\u6258\uff0cUI\u5c42\u8ba2\u9605\u8fd9\u4e9b\u59d4\u6258\u540e\u53ef\u5373\u65f6\u540c\u6b65\u5237\u65b0\u80cc\u5305\u754c\u9762\u3002"
    ))

    elems.append(make_heading3(None, "5.3.5.3\u3000\u7269\u54c1\u62fe\u53d6\u6d41\u7a0b"))
    elems.append(make_body(
        "\u573a\u666f\u4e2d\u7684\u53ef\u62fe\u53d6\u7269\u54c1\u7531\u9644\u52a0\u4e86UItemComponent\u7684\u666e\u901a"
        "Actor\u8868\u793a\uff0cUItemComponent\u6301\u6709\u5b8c\u6574\u7684FItemManifest\u914d\u7f6e\u3002"
        "\u73a9\u5bb6\u6309\u4e0b\u4ea4\u4e92\u952e\u65f6\uff0cAInventoryPlayerController::PrimaryInteract()"
        "\u68c0\u6d4b\u8303\u56f4\u5185\u5b9e\u73b0\u4e86IHighlightable\u63a5\u53e3\u7684\u76ee\u6807\u540e"
        "\u8c03\u7528UItemComponent::PickUp()\uff0c\u9500\u6bc1\u7269\u54c1Actor\u5e76\u5c06ItemComponent"
        "\u4f20\u5165UInventoryComp::TryAddItem()\u3002"
        "TryAddItem()\u5148\u901a\u8fc7HasRoomForItem()\u68c0\u67e5\u80cc\u5305\u5269\u4f59\u7a7a\u95f4\uff0c"
        "\u82e5\u7a7a\u95f4\u5145\u8db3\u5219\u901a\u8fc7Server RPC\u8f6c\u81f3\u670d\u52a1\u7aef\u6267\u884cServerAddNewItem()\uff0c"
        "\u5728FInventoryFastArray\u4e2d\u521b\u5efa\u65b0\u7684UInventoryItem\u5b9e\u4f8b\u5e76\u8c03\u7528MarkItemDirty()\u89e6\u53d1\u7f51\u7edc\u540c\u6b65\u3002"
    ))

    elems.append(make_heading3(None, "5.3.5.4\u3000\u6d88\u8017\u7cfb\u7edf\u4e0eGAS\u96c6\u6210"))
    elems.append(make_body(
        "Server_ConsumeItem()\u5728\u670d\u52a1\u7aef\u9012\u51cf\u7269\u54c1\u7684StackCount\uff1b"
        "\u5f53StackCount\u5f52\u96f6\u65f6\u8c03\u7528RemoveEntry()\u79fb\u9664\u8be5\u6761\u76ee\u3002"
        "\u6d88\u8017\u903b\u8f91\u7531FConsumableFragment::OnConsume()\u9a71\u52a8\uff0c"
        "\u5176\u5185\u90e8\u904d\u5386\u6240\u6709ConsumeModifier\uff0c\u6bcf\u4e2a\u5177\u4f53Modifier"
        "\uff08\u5982FHealthFragment\uff09\u8c03\u7528ApplyGameplayEffectByValue()\u5c06\u6548\u679c\u5e94\u7528\u81f3\u73a9\u5bb6ASC\uff1a"
        "\u51fd\u6570\u5185\u5148\u901a\u8fc7MakeOutgoingSpec()\u521b\u5efa\u9884\u914d\u7f6e\u7684GameplayEffect\u89c4\u8303\uff0c"
        "\u518d\u901a\u8fc7AssignTagSetByCallerMagnitude()\u4ee5SetByCaller\u65b9\u5f0f"
        "\u5c06\u53c2\u6570\u5316\u6570\u503c\u4e0eModifyTag\u5173\u8054\uff0c\u6700\u540e\u8c03\u7528"
        "ApplyGameplayEffectSpecToSelf()\u5b8c\u6210\u5c5e\u6027\u4fee\u6539\u3002"
        "\u8be5\u65b9\u5f0f\u4f7f\u5404\u79cd\u6d88\u8017\u7269\u54c1\u7684\u6548\u679c\u6570\u503c\u5728\u84dd\u56fe\u4e2d\u914d\u7f6e\uff0c\u65e0\u9700\u4fee\u6539C++\u4ee3\u7801\u3002"
    ))

    code_535 = [
        "void FEquipmentModifier::ApplyGameplayEffectByValue(",
        "    APlayerController* PC, float InValue)",
        "{",
        "    if (AARPGBaseCharacter* Char = Cast<AARPGBaseCharacter>(PC->GetPawn()))",
        "    {",
        "        if (UAbilitySystemComponent* ASC = Char->GetAbilitySystemComponent())",
        "        {",
        "            auto EffectCtx = ASC->MakeEffectContext();",
        "            FGameplayEffectSpecHandle Spec = ASC->MakeOutgoingSpec(",
        "                ModifyEffectClass, 1, EffectCtx);",
        "            UAbilitySystemBlueprintLibrary::AssignTagSetByCallerMagnitude(",
        "                Spec, ModifyTag, InValue);",
        "            ASC->ApplyGameplayEffectSpecToSelf(*Spec.Data.Get());",
        "        }",
        "    }",
        "}",
    ]
    for line in code_535:
        elems.append(make_code(line))

    elems.append(make_heading3(None, "5.3.5.5\u3000\u88c5\u5907\u7cfb\u7edf\u4e0e\u4e8b\u4ef6\u603b\u7ebf"))
    elems.append(make_body(
        "\u88c5\u5907\u4ea4\u4e92\u901a\u8fc7Server_EquipSlotClicked() Server RPC"
        "\u8f6c\u81f3\u670d\u52a1\u7aef\u5904\u7406\uff0c\u518d\u7531Multicast_EquipSlotClicked()\u5e7f\u64ad\u81f3\u6240\u6709\u7aef\uff0c"
        "\u786e\u4fdd\u73a9\u5bb6\u672c\u5730\u4e0e\u8fdc\u7aef\u5ba2\u6237\u7aef\u7684\u88c5\u5907\u72b6\u6001\u4fdd\u6301\u4e00\u81f4\u3002"
        "\u88c5\u5907\u903b\u8f91\u7684\u6838\u5fc3\u5728FEquipmentFragment::OnEquip()\u4e2d\uff1a"
        "\u51fd\u6570\u901a\u8fc7ARPG_EVENT_WITH_UOBJECT_TARGET\u5b8f\u5411\u89d2\u8272\u5e7f\u64ad"
        "\u201cCharacterEquipWeapon\u201d\u4e8b\u4ef6\uff0c\u4e8b\u4ef6\u6570\u636eUARPGEventData_Weapon"
        "\u643a\u5e26\u6b66\u5668\u7c7b\u578b\u3001\u914d\u7f6eID\u4e0e\u88c5\u5907Actor\u5f15\u7528\uff1b"
        "\u89d2\u8272\u6536\u5230\u4e8b\u4ef6\u540e\u66f4\u65b0\u653b\u51fb\u914d\u7f6e\u5e76\u89e6\u53d1\u6362\u6b66\u52a8\u753b\u3002"
        "\u540c\u65f6\uff0cUInventoryEquipmentComponent\u54cd\u5e94OnItemEquipped\u59d4\u6258\uff0c"
        "\u8c03\u7528SpawnEquippedActor()\u5728\u9aa8\u9abc Socket\u4f4d\u7f6e\u751f\u6210\u88c5\u5907Actor\uff0c"
        "\u5e76\u8c03\u7528\u5404EquipmentModifier\u7684OnEquip()\u901a\u8fc7GAS\u5e94\u7528\u88c5\u5907\u5e26\u6765\u7684\u5c5e\u6027\u52a0\u6210\u3002"
    ))

    # Sequence diagram
    elems.append(make_body(
        "\u5c40\u5185\u7269\u54c1\u7cfb\u7edf\u5404\u6d41\u7a0b\u7684\u5b8c\u6574\u8c03\u7528\u65f6\u5e8f\u5982\u56fe5-12\u6240\u793a\u3002"
    ))
    prefix_535 = "5_3_5_sequence"
    cap_535    = "\u56fe5-12\u2002\u5c40\u5185\u7269\u54c1\u7cfb\u7edf\u65f6\u5e8f\u56fe"
    img_path_535 = UML_DIR / f"{prefix_535}.png"
    if img_path_535.exists() and prefix_535 in downloaded:
        elems.append(make_image(img_path_535))
    else:
        elems.append(make_placeholder(f"{prefix_535}.png"))
    elems.append(make_caption(cap_535))

    return elems

# ──────────────────────────────────────────
# Insert into thesis
# ──────────────────────────────────────────
print("\n" + "=" * 55)
print("Step 2: Modify thesis docx")
print("=" * 55)

doc = Document(str(V2))

def find_para(doc, keyword):
    for p in doc.paragraphs:
        if keyword in p.text:
            return p
    return None

# -- 5.3.4 + 5.3.5: insert before 5.4 UI section --
ui_para = find_para(doc, "UI\u7cfb\u7edf\u5b9e\u73b0")
if ui_para is None:
    print("  Warning: could not find '5.4 UI system' heading, appending to end")
    for elem in build_534() + build_535():
        doc.add_paragraph()._element.getparent().append(elem)
else:
    ref = ui_para._element
    for elem in build_534() + build_535():
        ref.addprevious(elem)
    print("  Inserted 5.3.4 and 5.3.5 before '5.4 UI system'")

# -- 5.2.1 character select: insert sequence diagram at end of section --
cs_para = find_para(doc, "\u89d2\u8272\u9009\u62e9\u5b9e\u73b0")
shop_para = find_para(doc, "\u5546\u5e97\u5b9e\u73b0")
if cs_para and shop_para:
    ref = shop_para._element
    intro_521 = make_body(
        "\u89d2\u8272\u9009\u62e9\u6a21\u5757\u5404\u7ec4\u4ef6\u7684\u5b8c\u6574\u4ea4\u4e92\u65f6\u5e8f\u5982\u56fe5-9\u6240\u793a\u3002"
    )
    ref.addprevious(intro_521)
    cap_521  = "\u56fe5-9\u2002\u89d2\u8272\u9009\u62e9\u6a21\u5757\u65f6\u5e8f\u56fe"
    cap_elem = make_caption(cap_521)
    img_path_521 = UML_DIR / "5_2_1_sequence.png"
    if img_path_521.exists() and "5_2_1_sequence" in downloaded:
        img_elem = make_image(img_path_521)
    else:
        img_elem = make_placeholder("5_2_1_sequence.png")
    ref.addprevious(img_elem)
    ref.addprevious(cap_elem)
    cap_elem.addprevious(img_elem)
    print("  Inserted seq diagram 5-9 into 5.2.1")

# -- 5.2.2 shop: insert before next section heading (combat) --
combat_para = find_para(doc, "\u5c40\u5185\u6218\u6597\u6a21\u5757\u5b9e\u73b0")
if shop_para and combat_para:
    ref = combat_para._element
    intro_522 = make_body(
        "\u5546\u5e97\u6a21\u5757\u5404\u6d41\u7a0b\u7684\u5b8c\u6574\u65f6\u5e8f\u5982\u56fe5-10\u6240\u793a\u3002"
    )
    ref.addprevious(intro_522)
    cap_522  = "\u56fe5-10\u2002\u5546\u5e97\u6a21\u5757\u65f6\u5e8f\u56fe"
    cap_elem2 = make_caption(cap_522)
    img_path_522 = UML_DIR / "5_2_2_sequence.png"
    if img_path_522.exists() and "5_2_2_sequence" in downloaded:
        img_elem2 = make_image(img_path_522)
    else:
        img_elem2 = make_placeholder("5_2_2_sequence.png")
    ref.addprevious(img_elem2)
    ref.addprevious(cap_elem2)
    cap_elem2.addprevious(img_elem2)
    print("  Inserted seq diagram 5-10 into 5.2.2")

doc.save(str(V3))
print(f"\nSaved: {V3}")

# ──────────────────────────────────────────
# Update appendix
# ──────────────────────────────────────────
print("\n" + "=" * 55)
print("Step 3: Update appendix docx")
print("=" * 55)

NEW_APPENDIX_ENTRIES = [
    ("\u56fe5-9\u2002\u89d2\u8272\u9009\u62e9\u6a21\u5757\u65f6\u5e8f\u56fe\uff085.2.1\uff09",
     "5_2_1_sequence", SEQ_DIAGRAMS[0][3]),
    ("\u56fe5-10\u2002\u5546\u5e97\u6a21\u5757\u65f6\u5e8f\u56fe\uff085.2.2\uff09",
     "5_2_2_sequence", SEQ_DIAGRAMS[1][3]),
    ("\u56fe5-11\u2002\u4f24\u5bb3\u7cfb\u7edf\u65f6\u5e8f\u56fe\uff085.3.4\uff09",
     "5_3_4_sequence", SEQ_DIAGRAMS[2][3]),
    ("\u56fe5-12\u2002\u5c40\u5185\u7269\u54c1\u7cfb\u7edf\u65f6\u5e8f\u56fe\uff085.3.5\uff09",
     "5_3_5_sequence", SEQ_DIAGRAMS[3][3]),
]

app = Document(str(APP_PATH)) if APP_PATH.exists() else Document()
if not APP_PATH.exists():
    for section in app.sections:
        section.top_margin = Mm(30); section.bottom_margin = Mm(25)
        section.left_margin = Mm(30); section.right_margin = Mm(20)

app.add_page_break()

hdr = app.add_paragraph()
hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_hdr = hdr.add_run("\u7b2c\u4e94\u7ae0\u65f6\u5e8f\u56fe\u6e90\u7801\uff08\u65b0\u589e\uff09")
set_run_font(r_hdr, cn='\u9ed1\u4f53', en='Times New Roman', size=Pt(14), bold=True)

for caption, prefix, puml in NEW_APPENDIX_ENTRIES:
    ph = app.add_paragraph()
    ph.paragraph_format.space_before = Pt(12)
    rh = ph.add_run(caption)
    set_run_font(rh, cn='\u9ed1\u4f53', en='Times New Roman', size=Pt(13), bold=True)

    img_path = UML_DIR / f"{prefix}.png"
    if img_path.exists():
        pi = app.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi.add_run().add_picture(str(img_path), width=Inches(5.5))
        pc = app.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rpc = pc.add_run(caption)
        set_run_font(rpc, cn='\u5b8b\u4f53', size=SZ_5)

    src_title = app.add_paragraph()
    rs = src_title.add_run("PlantUML\u6e90\u7801\uff1a")
    set_run_font(rs, cn='\u9ed1\u4f53', size=Pt(11), bold=True)

    for line in puml.strip().split('\n'):
        pc2 = app.add_paragraph()
        pc2.paragraph_format.left_indent  = Pt(24)
        pc2.paragraph_format.space_before = Pt(0)
        pc2.paragraph_format.space_after  = Pt(0)
        pc2.paragraph_format.line_spacing = Pt(16)
        rcode = pc2.add_run(line)
        rcode.font.name = 'Courier New'
        rcode.font.size = Pt(9)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pc2._p.pPr.append(shd)

    app.add_paragraph()

app.save(str(APP_PATH))
print(f"Appendix saved: {APP_PATH}")

print("\n" + "=" * 55)
print("Done!")
print(f"  Thesis v3  : {V3}")
print(f"  Appendix   : {APP_PATH}")
print(f"  PNG folder : {UML_DIR}")
print("=" * 55)
